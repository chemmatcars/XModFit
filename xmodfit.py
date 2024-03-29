import importlib

from PyQt5.QtWidgets import QWidget, QApplication, QPushButton, QLabel, QLineEdit, QVBoxLayout, QMessageBox, QCheckBox, \
    QComboBox, QListWidget, QDialog, QFileDialog, QAbstractItemView, QSplitter, QSizePolicy, QAbstractScrollArea, QHBoxLayout, QTextEdit, QShortcut,\
    QProgressDialog, QDesktopWidget, QSlider, QTabWidget, QMenuBar, QAction, QTableWidgetSelectionRange, QProgressBar, QMenu, QTableWidgetItem, QTreeWidgetItem
from PyQt5.QtGui import QKeySequence, QFont, QDoubleValidator, QIntValidator, QTextCursor
from PyQt5.QtCore import Qt, QProcess
from PyQt5 import uic
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.backends.backend_qt5agg import NavigationToolbar2QT as NavigationToolbar
from matplotlib.figure import Figure
import webbrowser, shutil
from docx import Document
import os
import glob
import sys
import pyqtgraph as pg
from pyqtgraph.dockarea import DockArea, Dock
from PlotWidget import PlotWidget
import copy
import numpy as np
from Data_Dialog import Data_Dialog
# from readData import read1DSAXS
from importlib import import_module, reload
from Fit_Routines import Fit
from tabulate import tabulate
import corner
import numbers
import time
import shutil
from FunctionEditor import FunctionEditor
from MultiInputDialog import MultiInputDialog
import traceback
import pandas as pd
from functools import partial
import pylab as pl
from scipy.stats import chi2
from scipy.interpolate import interp1d
import math
from mplWidget import MplWidget
import statsmodels.api as sm
import Chemical_Formula
import Structure_Factors
import utils
import xraydb
from collections import OrderedDict




class minMaxDialog(QDialog):
    def __init__(self, value, vary=0, minimum=None, maximum=None, expr=None, brute_step=None, parent=None, title=None):
        QDialog.__init__(self, parent)
        self.value = value
        self.vary = vary
        if minimum is None:
            self.minimum = -np.inf
        else:
            self.minimum = minimum
        if maximum is None:
            self.maximum = np.inf
        else:
            self.maximum = maximum
        self.expr = expr
        self.brute_step = brute_step
        self.createUI()
        if title is not None:
            self.setWindowTitle(title)
        
    def createUI(self):
        self.vblayout = QVBoxLayout(self)
        self.layoutWidget = pg.LayoutWidget()
        self.vblayout.addWidget(self.layoutWidget)
        
        valueLabel = QLabel('Value')
        self.layoutWidget.addWidget(valueLabel)
        self.layoutWidget.nextColumn()
        self.valueLineEdit = QLineEdit(str(self.value))
        self.layoutWidget.addWidget(self.valueLineEdit)

        self.layoutWidget.nextRow()
        varyLabel = QLabel('Fit')
        self.layoutWidget.addWidget(varyLabel)
        self.layoutWidget.nextColumn()
        self.varyCheckBox = QCheckBox()
        self.layoutWidget.addWidget(self.varyCheckBox)
        if self.vary>0:
            self.varyCheckBox.setCheckState(Qt.Checked)
        else:
            self.varyCheckBox.setCheckState(Qt.Unchecked)

        self.layoutWidget.nextRow()
        minLabel = QLabel('Minimum')
        self.layoutWidget.addWidget(minLabel)
        self.layoutWidget.nextColumn()
        self.minimumLineEdit = QLineEdit(str(self.minimum))
        self.layoutWidget.addWidget(self.minimumLineEdit)
        
        self.layoutWidget.nextRow()
        maxLabel = QLabel('Maximum')
        self.layoutWidget.addWidget(maxLabel)
        self.layoutWidget.nextColumn()
        self.maximumLineEdit = QLineEdit(str(self.maximum))
        self.layoutWidget.addWidget(self.maximumLineEdit)
        
        self.layoutWidget.nextRow()
        exprLabel = QLabel('Expr')
        self.layoutWidget.addWidget(exprLabel)
        self.layoutWidget.nextColumn()
        self.exprLineEdit = QLineEdit(str(self.expr))
        self.layoutWidget.addWidget(self.exprLineEdit)
        
        self.layoutWidget.nextRow()
        bruteStepLabel = QLabel('Brute step')
        self.layoutWidget.addWidget(bruteStepLabel)
        self.layoutWidget.nextColumn()
        self.bruteStepLineEdit = QLineEdit(str(self.brute_step))
        self.layoutWidget.addWidget(self.bruteStepLineEdit)
        
        self.layoutWidget.nextRow()
        self.cancelButton = QPushButton('Cancel')
        self.cancelButton.clicked.connect(self.cancelandClose)
        self.layoutWidget.addWidget(self.cancelButton)
        self.layoutWidget.nextColumn()
        self.okButton = QPushButton('OK')
        self.okButton.clicked.connect(self.okandClose)
        self.layoutWidget.addWidget(self.okButton)
        self.okButton.setDefault(True)
        
    def okandClose(self):
        # try:
        if type(eval(self.valueLineEdit.text())*1.0)==float:
            self.value = float(self.valueLineEdit.text())
        else:
            QMessageBox.warning(self, 'Value Error',
                                'Please enter floating point number for Value', QMessageBox.Ok)
            self.minimumLineEdit.setText(str(self.minimum))
            return
        if self.varyCheckBox.checkState() == Qt.Checked:
            self.vary = 1
        else:
            self.vary = 0

        minimum=self.minimumLineEdit.text()
        if '-inf' in minimum:
            self.minimum=-np.inf
        elif type(eval(self.minimumLineEdit.text())*1.0)==float:
            self.minimum=float(self.minimumLineEdit.text())
        else:
            QMessageBox.warning(self,'Value Error',
                                'Please enter floating point number for Minimum value',QMessageBox.Ok)
            self.minimumLineEdit.setText(str(self.minimum))
            return

        maximum = self.maximumLineEdit.text()
        if 'inf' in maximum:
            self.maximum=np.inf
        elif type(eval(self.maximumLineEdit.text())*1.0)==float:
            self.maximum = float(self.maximumLineEdit.text())
        else:
            QMessageBox.warning(self, 'Value Error',
                                'Please enter floating point number for Maximum value', QMessageBox.Ok)
            self.maximumLineEdit.setText(str(self.maximum))
            return

        self.expr=self.exprLineEdit.text()
        if self.expr != 'None':
            self.vary=0
        if self.bruteStepLineEdit.text() != 'None':
            self.brute_step = float(self.bruteStepLineEdit.text())
        else:
            self.brute_step = None
        self.accept()
        # except:
        #     QMessageBox.warning(self,'Value Error','Value, Min, Max should be floating point numbers\n\n'+traceback.format_exc(),QMessageBox.Ok)

    def cancelandClose(self):
        self.reject()

class FitResultDialog(QDialog):
    def __init__(self,fit_report,fit_info,parent=None):
        QDialog.__init__(self,parent)
        self.setWindowTitle('Fit Results')
        self.fit_report=fit_report
        self.fit_info=fit_info
        self.createUI()
        self.resize(600,400)
        
    def createUI(self):
        self.vblayout=QVBoxLayout(self)
        self.layoutWidget=pg.LayoutWidget()
        self.vblayout.addWidget(self.layoutWidget)
        
        fitReportLabel=QLabel('Fit Report')
        self.layoutWidget.addWidget(fitReportLabel,colspan=2)
        self.layoutWidget.nextRow()
        self.fitReportTextEdit=QTextEdit()
        self.fitReportTextEdit.setText(self.fit_report)
        self.layoutWidget.addWidget(self.fitReportTextEdit,colspan=2)
        
        self.layoutWidget.nextRow()
        fitInfoLabel=QLabel('Fit Info')
        self.layoutWidget.addWidget(fitInfoLabel,colspan=2)
        self.layoutWidget.nextRow()
        self.fitInfoTextEdit=QTextEdit()
        self.fitInfoTextEdit.setText(self.fit_info)
        self.layoutWidget.addWidget(self.fitInfoTextEdit,colspan=2)
        
        self.layoutWidget.nextRow()
        self.cancelButton=QPushButton('Reject')
        self.cancelButton.clicked.connect(self.cancelandClose)
        self.layoutWidget.addWidget(self.cancelButton,col=0)
        self.okButton=QPushButton('Accept')
        self.okButton.clicked.connect(self.okandClose)
        self.layoutWidget.addWidget(self.okButton,col=1)
        self.okButton.setDefault(True)
        
    def okandClose(self):
        self.accept()
        
    def cancelandClose(self):
        self.reject()


class XModFit(QWidget):
    """
    This widget class is developed to provide an end-user a *Graphical User Interface* by which either they can \
    develop their own fitting functions in python or use the existing fitting functions under different categories\
     to analyze different kinds of one-dimensional data sets. `LMFIT <https://lmfit.github.io/lmfit-py/>`_ is extensively\
      used within this widget.
    
    **Features**
    
    1. Read and fit multiple data files
    2. Already available functions are categorized as per the function types and techniques
    3. Easy to add more catergories and user-defined functions
    4. Once the function is defined properly all the free and fitting parameters will be available within the GUI as tables.
    5. An in-built Function editor with a function template is provided.
    6. The function editor is enabled with python syntax highlighting.
    
    **Usage**
    
    :class:`Fit_Widget` can be used as stand-alone python fitting package by running it in terminal as::
    
        $python xmodfit.py
        
    .. figure:: Figures/Fit_widget.png
       :figwidth: 100%
       
       **Fit Widget** in action.
    
    Also it can be used as a widget with any other python application.
    """
    
    def __init__(self,parent=None):
        QWidget.__init__(self,parent)
        self.vblayout=QVBoxLayout(self)
        self.menuBar = QMenuBar(self)
        self.menuBar.setNativeMenuBar(False)
        self.create_menus()
        self.vblayout.addWidget(self.menuBar,0)
        self.mainDock=DockArea(self,parent)
        self.vblayout.addWidget(self.mainDock,5)

        self.emcee_walker = 100
        self.emcee_steps = 100
        self.emcee_burn = 0
        self.emcee_thin = 1
        self.emcee_cores = 1
        self.emcee_frac = self.emcee_burn/self.emcee_steps
        self.reuse_sampler = False
        self.autoCalculate = True
        self.funcDock=Dock('Functions',size=(1,6),closable=False,hideTitle=False)
        self.fitDock=Dock('Fit options',size=(1,2),closable=False,hideTitle=False)
        self.dataDock=Dock('Data',size=(1,8),closable=False,hideTitle=False)
        self.paramDock=Dock('Parameters',size=(2,8),closable=False,hideTitle=False)
        self.plotDock=Dock('Data and Fit',size=(5,8),closable=False,hideTitle=False)
        self.fitResultDock=Dock('Fit Results',size=(5,8),closable=False,hideTitle=False)
        self.mainDock.addDock(self.dataDock)
        self.mainDock.addDock(self.fitDock,'bottom')
        self.mainDock.addDock(self.paramDock,'right')
        self.mainDock.addDock(self.fitResultDock, 'right')
        self.mainDock.addDock(self.plotDock,'above',self.fitResultDock)

        self.mainDock.addDock(self.funcDock,'above',self.dataDock)
        self.special_keys=['x','params','choices','output_params','__mpar__']
        self.curr_funcClass={}
        
        
        self.data={}
        self.dlg_data={}
        self.plotColIndex={}
        self.plotColors={}
        self.curDir=os.getcwd()
        self.fileNumber=0
        self.fileNames={}
        self.fchanged=True
        self.chisqr='None'
        self.format='%.6e'
        self.gen_param_items=[]
        self.doubleValidator=QDoubleValidator()
        self.intValidator=QIntValidator()
        self.tApp_Clients={}
        self.tModules={}
        self.fitMethods={'Levenberg-Marquardt':'leastsq',
                         'Scipy-Least-Squares':'least_squares',
                         'Differential-Evolution': 'differential_evolution'}
                         # 'Brute-Force-Method':'brute',
                         # 'Nelder-Mead':'nelder',
                         # 'L-BFGS-B':'lbfgsb',
                         # 'Powell':'powell',
                         # 'Congugate-Gradient':'cg',
                         # 'Newton-CG-Trust-Region':'trust-ncg',
                         # 'COBLYA':'cobyla',
                         # 'Truncate-Newton':'tnc',
                         # 'Exact-Trust-Region':'trust-exact',
                         # 'Dogleg':'dogleg',
                         # 'Sequential-Linear-Square':'slsqp',
                         # 'Adaptive-Memory-Programming':'ampgo',
                         # 'Maximum-Likelihood-MC-Markov-Chain':'emcee'}
                         #
        
        self.create_funcDock()
        self.create_fitDock()
        self.create_dataDock()
        self.create_plotDock()
        self.create_fitResultDock()
        self.update_catagories()
        self.create_paramDock()
        # self.xminmaxChanged()
        self.sfnames=None
        self.expressions={}

    def create_menus(self):
        self.fileMenu = self.menuBar.addMenu('&File')
        self.settingsMenu = self.menuBar.addMenu('&Settings')
        self.toolMenu = self.menuBar.addMenu('&Tools')
        self.helpMenu = self.menuBar.addMenu('&Help')

        quit=QAction('Quit',self)
        quit.triggered.connect(self.close)
        self.fileMenu.addAction(quit)

        parFormat=QAction('&Parameter format',self)
        parFormat.triggered.connect(self.changeParFormat)
        self.settingsMenu.addAction(parFormat)

        about=QAction('&About',self)
        about.triggered.connect(self.aboutDialog)
        self.helpMenu.addAction(about)


        toolItems=os.listdir(os.path.join(os.curdir,'Tools'))
        self.toolDirs=[]
        self.toolApps={}
        for item in toolItems:
            if '__' not in item:
                self.toolDirs.append(self.toolMenu.addMenu('&%s'%item))
                tApps=glob.glob(os.path.join(os.curdir,'Tools',item,'*.py'))
                for app in tApps:
                    tname='&'+os.path.basename(os.path.splitext(app)[0])
                    self.toolApps[tname]=app
                    tApp=QAction(tname,self)
                    tApp.triggered.connect(self.launch_tApp)
                    self.toolDirs[-1].addAction(tApp)

    def changeParFormat(self):
        dlg=MultiInputDialog(inputs={'Format':self.format},title='Parameter format')
        if dlg.exec_():
            self.format=dlg.inputs['Format']
            try:
                self.update_sfit_parameters()
                self.update_mfit_parameters_new()
            except:
                pass

    def launch_tApp(self):
        tname=self.sender().text()
        module_name=".".join(os.path.splitext(self.toolApps[tname])[0].split(os.sep)[1:])
        if module_name not in sys.modules:
            self.tModules[module_name]=importlib.import_module(module_name)
        tmodule=self.tModules[module_name]
        if tmodule in self.tApp_Clients:
            self.tApp_Clients[tmodule].show()
        else:
            tclass = getattr(tmodule, tname[1:])
            self.tApp_Clients[tmodule]=tclass(self)
            self.tApp_Clients[tmodule].setWindowTitle(tname[1:])
            self.tApp_Clients[tmodule].show()








        # if tname not in self.tApp_Clients or self.tApp_Clients[tname].pid() is None:
        #     self.tApp_Clients[tname]=QProcess()
        #     self.tApp_Clients[tname].start('python '+self.toolApps[tname])
        # elif self.tApp_Clients[tname].pid()>0:
        #     QMessageBox.warning(self,'Running...','The tool %s is already running'%tname,QMessageBox.Ok)
        # else:
        #     self.tApp_Clients[tname].start('python ' + self.toolApps[tname])



    def aboutDialog(self):
        QMessageBox.information(self,'About','Copyright (c) 2022 NSF\'s ChemMAtCARS, University of Chicago.\n\n'
                                             'Developers:\n'
                                             'Mrinal K. Bera (mrinalkb@cars.uchicago.edu \n'
                                             'Wei Bu (bu@cars.uchicago.edu)'
                                             ,QMessageBox.Ok)
        
    def create_funcDock(self):
        self.funcLayoutWidget=pg.LayoutWidget(self)
        row=0
        col=0
        funcCategoryLabel=QLabel('Function Categories:')
        self.funcLayoutWidget.addWidget(funcCategoryLabel,row=row,col=col,colspan=2)
        
        row+=1
        col=0
        self.addCategoryButton=QPushButton('Create')
        self.addCategoryButton.clicked.connect(self.addCategory)
        self.funcLayoutWidget.addWidget(self.addCategoryButton,row=row,col=col)
        col+=1
        self.removeCategoryButton=QPushButton('Remove')
        self.removeCategoryButton.clicked.connect(self.removeCategory)
        self.funcLayoutWidget.addWidget(self.removeCategoryButton,row=row,col=col)
        
        row+=1
        col=0
        self.categoryListWidget=QListWidget()
        self.categoryListWidget.currentItemChanged.connect(self.update_functions)
        self.funcLayoutWidget.addWidget(self.categoryListWidget,row=row,col=col,colspan=2)
        
        row+=1
        col=0
        funcLabel=QLabel('Functions:')
        self.funcLayoutWidget.addWidget(funcLabel,row=row,col=col,colspan=2)
        
        row+=1
        col=0
        self.addFuncButton=QPushButton('Create')
        self.addFuncButton.clicked.connect(self.addFunction)
        self.funcLayoutWidget.addWidget(self.addFuncButton,row=row,col=col)
        col+=1
        self.removeFuncButton=QPushButton('Remove')
        self.removeFuncButton.clicked.connect(self.removeFunction)
        self.funcLayoutWidget.addWidget(self.removeFuncButton,row=row,col=col)
        
        row+=1
        col=0
        self.funcListWidget=QListWidget()
        self.funcListWidget.setSelectionMode(4)
        self.funcListWidget.setContextMenuPolicy(Qt.CustomContextMenu)
        self.funcListWidget.customContextMenuRequested.connect(self.funcListRightClicked)
        self.funcListWidget.itemSelectionChanged.connect(self.functionChanged)
        self.funcListWidget.itemDoubleClicked.connect(self.openFunction)
        self.funcLayoutWidget.addWidget(self.funcListWidget,row=row,col=col,colspan=2)
        
        self.funcDock.addWidget(self.funcLayoutWidget)

    def funcListRightClicked(self,pos):
        popMenu = QMenu()
        showDet = QAction("Show Details", self)
        addDet = QAction("Upload Details", self)
        modDet = QAction("Create/Modify Details", self)
        popMenu.addAction(showDet)
        popMenu.addAction(addDet)
        popMenu.addAction(modDet)
        showDet.triggered.connect(self.showDetails)
        addDet.triggered.connect(self.addDetails)
        modDet.triggered.connect(self.modifyDetails)
        popMenu.exec_(self.funcListWidget.mapToGlobal(pos))

    def showDetails(self):
        url = os.path.join(os.path.curdir, 'Function_Details', self.categoryListWidget.currentItem().text(),
                            self.funcListWidget.currentItem().text(),'help.pdf')
        if os.path.exists(url):
            webbrowser.open_new_tab(url)
        else:
            QMessageBox.warning(self,'File Error','The help files regarding the function details do not exist.',QMessageBox.Ok)
        # os.system('C:/Users/mrinalkb/Desktop/ESH738.pdf')

    def addDetails(self):
        path=os.path.join(os.path.curdir,'Function_Details',self.categoryListWidget.currentItem().text(),self.funcListWidget.currentItem().text())
        if os.path.exists(path):
            fname = QFileDialog.getOpenFileName(self,caption='Select help file',directory=self.curDir,filter="Help files (*.docx *.pdf)")[0]
            tfname=os.path.join(path,'help'+os.path.splitext(fname)[1])
            shutil.copy(fname,tfname)
        else:
            os.makedirs(path)

    def modifyDetails(self):
        category=self.categoryListWidget.currentItem().text()
        function=self.funcListWidget.currentItem().text()
        path = os.path.join(os.path.curdir, 'Function_Details', category,
                            function,'help.docx')
        if os.path.exists(path):
            webbrowser.open_new_tab(path)
        else:
            if not os.path.exists(os.path.dirname(path)):
                os.makedirs(os.path.dirname(path))
            doc=Document()
            doc.add_heading('Details of %s/%s'%(category,function),0)
            module = 'Functions.%s.%s' % (category,function)
            text=getattr(self.curr_funcClass[module], function).__init__.__doc__
            doc.add_paragraph(text)
            doc.save(path)
            webbrowser.open_new_tab(path)

        
    def addCategory(self):
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
        tdir=QFileDialog.getExistingDirectory(self,'Select a folder','./Functions/',QFileDialog.ShowDirsOnly)
        if tdir!='': 
            cdir=os.path.basename(os.path.normpath(tdir))
            fh=open(os.path.join(tdir,'__init__.py'),'w')
            fh.write('__all__=[]')
            fh.close()
            if cdir not in self.categories:
                self.categories.append(cdir)
                self.categoryListWidget.addItem(cdir)
            else:
                QMessageBox.warning(self,'Category error','Category already exist!',QMessageBox.Ok)

        
    def removeCategory(self):
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
        self.funcListWidget.clear()
        if len(self.categoryListWidget.selectedItems())==1:
            ans=QMessageBox.question(self,'Delete warning','Are you sure you would like to delete the category?',
                                     QMessageBox.No,QMessageBox.Yes)
            if ans==QMessageBox.Yes:
                category=os.path.abspath('./Functions/%s'%self.categoryListWidget.currentItem().text())
                #os.rename(category,)
                shutil.rmtree(category)
                self.categories.remove(self.categoryListWidget.currentItem().text())
                self.categoryListWidget.takeItem(self.categoryListWidget.currentRow())
        elif len(self.categoryListWidget.selectedItems())>1:
            QMessageBox.warning(self,'Warning','Please select only one category at a time to remove',QMessageBox.Ok)
        else:
            QMessageBox.warning(self,'Warning','Please select one category atleast to remove',QMessageBox.Ok)
            
            
    def openFunction(self):
        dirName=os.path.abspath('./Functions/%s'%self.categoryListWidget.currentItem().text())
        funcName=self.funcListWidget.currentItem().text()
        try:
            if not self.funcEditor.open: 
                self.funcEditor=FunctionEditor(funcName=funcName,dirName=dirName)
                self.funcEditor.setWindowTitle('Function editor')
                self.funcEditor.show()
                self.funcOpen=self.funcEditor.open
                self.funcEditor.closeEditorButton.clicked.connect(self.postAddFunction)
            else:
                QMessageBox.warning(self,'Warning','You cannot edit two functions together',QMessageBox.Ok)
        except:
            self.funcEditor=FunctionEditor(funcName=funcName,dirName=dirName)
            self.funcEditor.setWindowTitle('Function editor')
            self.funcEditor.show()
            self.funcEditor.closeEditorButton.clicked.connect(self.postAddFunction)
                
    def addFunction(self):
        if len(self.categoryListWidget.selectedItems())==1:
            dirName=os.path.abspath('./Functions/%s'%self.categoryListWidget.currentItem().text())
            self.funcEditor=FunctionEditor(dirName=dirName)
            self.funcEditor.setWindowTitle('Function editor')
            self.funcEditor.show()
            self.funcEditor.closeEditorButton.clicked.connect(self.postAddFunction)
        else:
            QMessageBox.warning(self,'Category Error','Please select a Category first',QMessageBox.Ok)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

        
        
    def postAddFunction(self):
        if self.funcEditor.funcNameLineEdit.text()!='tmpxyz':
            dirName=os.path.abspath('./Functions/%s'%self.categoryListWidget.currentItem().text())
            fh=open(os.path.join(dirName,'__init__.py'),'r')
            line=fh.readlines()
            fh.close()
            funcList=eval(line[0].split('=')[1])
            funcName=self.funcEditor.funcNameLineEdit.text()
            if funcName not in funcList:
                funcList.append(funcName)
                funcList=sorted(list(set(funcList)),key=str.lower)
                os.remove(os.path.join(dirName,'__init__.py'))
                fh=open(os.path.join(dirName,'__init__.py'),'w')
                fh.write('__all__='+str(funcList))
                fh.close()
            self.update_functions()
        
        
    
    def removeFunction(self):
        if len(self.funcListWidget.selectedItems())==1:
            ans=QMessageBox.question(self,'Warning','Are you sure you would like to remove the function',
                                     QMessageBox.No,QMessageBox.Yes)
            if ans==QMessageBox.Yes:
                dirName=os.path.abspath('./Functions/%s'%self.categoryListWidget.currentItem().text())
                fname=self.funcListWidget.currentItem().text()
                fh=open(os.path.join(dirName,'__init__.py'),'r')
                line=fh.readlines()
                fh.close()
                funcList=eval(line[0].split('=')[1])
                try:
                    os.remove(os.path.join(dirName,fname+'.py'))
                    os.remove(os.path.join(dirName,'__init__.py'))
                    fh=open(os.path.join(dirName,'__init__.py'),'w')
                    fh.write('__all__='+str(funcList))
                    fh.close()
                    self.update_functions()
                except:
                    QMessageBox.warning(self,'Remove error','Cannot remove the function because the function file\
                     might be open elsewhere.\n\n'+traceback.format_exc(),QMessageBox.Ok)
        elif len(self.funcListWidget.selectedItems())>1:
            QMessageBox.warning(self,'Warning','Please select only one function at a time to remove',QMessageBox.Ok)
        else:
            QMessageBox.warning(self,'Warning','Please select one function atleast to remove',QMessageBox.Ok)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
        
    def create_dataDock(self):
        self.dataLayoutWidget=pg.LayoutWidget(self)
        
        datafileLabel=QLabel('Data files')
        self.dataLayoutWidget.addWidget(datafileLabel,colspan=2)
        
        self.dataLayoutWidget.nextRow()
        self.addDataButton=QPushButton('Add files')
        self.dataLayoutWidget.addWidget(self.addDataButton)
        self.addDataButton.clicked.connect(lambda x: self.addData())
        self.removeDataButton=QPushButton('Remove Files')
        self.dataLayoutWidget.addWidget(self.removeDataButton,col=1)
        self.removeDataButton.clicked.connect(self.removeData)
        self.removeDataShortCut = QShortcut(QKeySequence.Delete, self)
        self.removeDataShortCut.activated.connect(self.removeData)
        
        
        self.dataLayoutWidget.nextRow()
        self.dataListWidget=QListWidget()
        self.dataListWidget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.dataListWidget.itemSelectionChanged.connect(self.dataFileSelectionChanged)
        self.dataListWidget.itemDoubleClicked.connect(self.openDataDialog)
        self.dataLayoutWidget.addWidget(self.dataListWidget,colspan=2)
        
        self.dataDock.addWidget(self.dataLayoutWidget)

    def create_fitDock(self):
        self.fitLayoutWidget=pg.LayoutWidget(self)

        xminmaxLabel = QLabel('Xmin:Xmax')
        self.fitLayoutWidget.addWidget(xminmaxLabel)
        self.xminmaxLineEdit = QLineEdit('0:1')
        self.xminmaxLineEdit.returnPressed.connect(self.xminmaxChanged)
        self.fitLayoutWidget.addWidget(self.xminmaxLineEdit, col=1)

        self.fitLayoutWidget.nextRow()
        fitMethodLabel = QLabel('Fit Method')
        self.fitLayoutWidget.addWidget(fitMethodLabel)
        self.fitMethodComboBox = QComboBox()
        self.fitMethodComboBox.addItems(list(self.fitMethods.keys()))
        self.fitLayoutWidget.addWidget(self.fitMethodComboBox, col=1)

        self.fitLayoutWidget.nextRow()
        fitScaleLabel = QLabel('Fit Scale')
        self.fitLayoutWidget.addWidget(fitScaleLabel)
        self.fitScaleComboBox = QComboBox()
        self.fitScaleComboBox.addItems(['Linear', 'Linear w/o error', 'Log', 'Log w/o error'])
        self.fitLayoutWidget.addWidget(self.fitScaleComboBox, col=1)

        self.fitLayoutWidget.nextRow()
        fitIterationLabel = QLabel('Fit Iterations')
        self.fitLayoutWidget.addWidget(fitIterationLabel)
        self.fitIterationLineEdit = QLineEdit('1000')
        self.fitLayoutWidget.addWidget(self.fitIterationLineEdit, col=1)

        self.fitLayoutWidget.nextRow()
        self.fitButton = QPushButton('Fit')
        self.fitButton.clicked.connect(lambda x: self.doFit())
        self.fitButton.setEnabled(False)
        self.unfitButton = QPushButton('Undo fit')
        self.unfitButton.clicked.connect(self.undoFit)
        self.fitLayoutWidget.addWidget(self.unfitButton)
        self.fitLayoutWidget.addWidget(self.fitButton, col=1)

        self.fitLayoutWidget.nextRow()
        confIntervalMethodLabel=QLabel('Confidence Interval Method')
        self.confIntervalMethodComboBox=QComboBox()
        self.confIntervalMethodComboBox.addItems(['ChiSqrDist', 'MCMC'])
        self.fitLayoutWidget.addWidget(confIntervalMethodLabel)
        self.fitLayoutWidget.addWidget(self.confIntervalMethodComboBox,col=1)

        self.fitLayoutWidget.nextRow()
        self.showConfIntervalButton = QPushButton('Show Param Error')
        self.showConfIntervalButton.setDisabled(True)
        self.showConfIntervalButton.clicked.connect(self.confInterval_emcee)
        self.calcConfInterButton = QPushButton('Calculate Param Error')
        self.calcConfInterButton.clicked.connect(self.calcConfInterval)
        self.calcConfInterButton.setDisabled(True)
        self.fitLayoutWidget.addWidget(self.showConfIntervalButton)
        self.fitLayoutWidget.addWidget(self.calcConfInterButton, col=1)

        self.fitDock.addWidget(self.fitLayoutWidget)


    def dataFileSelectionChanged(self):
        self.sfnames=[]
        self.pfnames=[]
        for item in self.dataListWidget.selectedItems():
            self.sfnames.append(item.text())
            txt=item.text()
            self.pfnames=self.pfnames+[txt.split('<>')[0]+':'+key for key in self.data[txt].keys()]
        if len(self.sfnames)>0:
            self.curDir = os.path.dirname(self.sfnames[-1].split('<>')[1])
            xmin=np.min([np.min([np.min(self.data[key][k1]['x']) for k1 in self.data[key].keys()]) for key in self.sfnames])
            xmax=np.max([np.max([np.max(self.data[key][k1]['x']) for k1 in self.data[key].keys()]) for key in self.sfnames])
            self.xminmaxLineEdit.setText('%0.3f:%0.3f'%(xmin,xmax))
            self.xminmaxChanged()
            # if len(self.data[self.sfnames[-1]].keys())>1:
            #     text='{'
            #     for key in self.data[self.sfnames[-1]].keys():
            #         text+='"'+key+'":np.linspace(%.3f,%.3f,%d),'%(xmin,xmax,100)
            #     text=text[:-1]+'}'
            # else:
            #     text='np.linspace(%.3f,%.3f,100)'%(xmin,xmax)
            # self.xLineEdit.setText(text)
            self.fitButton.setEnabled(True)
        else:
            self.fitButton.setDisabled(True)
            try:
                self.update_plot()
            except:
                pass
        # self.update_plot()
        # self.xChanged()
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
            
    def openDataDialog(self,item):
        fnum,fname=item.text().split('<>')
        self.dataListWidget.itemSelectionChanged.disconnect()
        data_dlg=Data_Dialog(data=self.dlg_data[item.text()],parent=self,expressions=self.expressions[item.text()],plotIndex=self.plotColIndex[item.text()],colors=self.plotColors[item.text()])
        data_dlg.setModal(True)
        data_dlg.closePushButton.setText('Cancel')
        data_dlg.tabWidget.setCurrentIndex(1)
        data_dlg.dataFileLineEdit.setText(fname)
        if data_dlg.exec_():
            self.plotWidget.remove_data(datanames=self.pfnames)
            newFname=data_dlg.dataFileLineEdit.text()
            if fname==newFname:
                self.plotColIndex[item.text()]=data_dlg.plotColIndex
                self.plotColors[item.text()]=data_dlg.plotColors
                self.dlg_data[item.text()]=copy.copy(data_dlg.data)
                self.data[item.text()]=copy.copy(data_dlg.externalData)
                self.expressions[item.text()]=data_dlg.expressions
                for key in self.data[item.text()].keys():
                    self.plotWidget.add_data(self.data[item.text()][key]['x'],self.data[item.text()][key]['y'],yerr=self.data[item.text()][key]['yerr'],name='%s:%s'%(fnum,key),color=self.plotColors[item.text()][key])
            else:
                text = '%s<>%s' % (fnum, newFname)
                self.data[text] = self.data.pop(item.text())
                self.dlg_data[text] = self.dlg_data.pop(item.text())
                item.setText(text)
                self.dlg_data[text]=copy.copy(data_dlg.data)
                self.data[text]=copy.copy(data_dlg.externalData)
                self.plotColIndex[text]=data_dlg.plotColIndex
                self.plotColors[text]=data_dlg.plotColors
                self.expressions[text]=data_dlg.expressions
                for key in self.data[text].keys():
                    self.plotWidget.add_data(self.data[text][key]['x'], self.data[text][key]['y'], yerr=self.data[text][key][
                    'yerr'],name='%s:%s'%(fnum,key),color=self.plotColors[text][key])
        # self.sfnames = []
        # self.pfnames = []
        # for item in self.dataListWidget.selectedItems():
        #     self.sfnames.append(item.text())
        #     txt=item.text()
        #     self.pfnames=self.pfnames+[txt.split('<>')[0]+':'+key for key in self.data[txt].keys()]
        self.dataFileSelectionChanged()
        # self.xChanged()
        self.dataListWidget.itemSelectionChanged.connect(self.dataFileSelectionChanged)
        #self.update_plot()

    def xminmaxChanged(self):
        try:
            xmin,xmax=self.xminmaxLineEdit.text().split(':')
            self.xmin, self.xmax=float(xmin),float(xmax)
            self.update_plot()
        except:
            QMessageBox.warning(self,"Value Error", "Please supply the Xrange in this format:\n xmin:xmax",QMessageBox.Ok)
    


    def doFit(self, fit_method=None, emcee_walker=100, emcee_steps=100,
                       emcee_cores=1, reuse_sampler=False, emcee_burn=30, emcee_thin=1):
        self.fchanged=False
        self.tchisqr=1e30
        self.xminmaxChanged()
        if self.sfnames is None or self.sfnames==[]:
            QMessageBox.warning(self,'Data Error','Please select a dataset first before fitting',QMessageBox.Ok)
            return
        try:
            if len(self.fit.fit_params)>0:
                pass
            else:
                QMessageBox.warning(self, 'Fit Warning', 'Please select atleast a single parameter to fit', QMessageBox.Ok)
                return
        except:
            QMessageBox.warning(self, 'Fit Function Warning', 'Please select a function to fit', QMessageBox.Ok)
            return
        if len(self.funcListWidget.selectedItems())==0:
            QMessageBox.warning(self, 'Function Error',
                                'Please select a function first to fit.\n' + traceback.format_exc(), QMessageBox.Ok)
            return
        # try:
        #     self.fixedParamTableWidget.cellChanged.disconnect(self.fixedParamChanged)
        #     self.sfitParamTableWidget.cellChanged.disconnect(self.sfitParamChanged)
        #     self.mfitParamTableWidget.cellChanged.disconnect(self.mfitParamChanged)
        # except:
        #     QMessageBox.warning(self,'Function Error','Please select a function first to fit.\n'+traceback.format_exc(),QMessageBox.Ok)
        #     return
        if fit_method is None:
            self.fit_method=self.fitMethods[self.fitMethodComboBox.currentText()]
        else:
            self.fit_method=fit_method
        if self.fit_method not in ['leastsq','brute','differential_evolution','least_squares','emcee']:
            QMessageBox.warning(self,'Fit Method Warning','This method is under development and will be available '
                                                          'soon. Please use only Lavenberg-Marquardt for the time '
                                                          'being.', QMessageBox.Ok)
            return
        self.fit_scale=self.fitScaleComboBox.currentText()
        try:
            self.fit.functionCalled.disconnect()
        except:
            pass
        if self.fit_method!='emcee':
            self.fit.functionCalled.connect(self.fitCallback)
        else:
            self.fit.functionCalled.connect(self.fitErrorCallback)
        for fname in self.sfnames:
            if len(self.data[fname].keys())>1:
                x={}
                y={}
                yerr={}
                for key in self.data[fname].keys():
                    x[key]=self.data[fname][key]['x']
                    y[key]=self.data[fname][key]['y']
                    yerr[key]=self.data[fname][key]['yerr']
            else:
                key=list(self.data[fname].keys())[0]
                x=self.data[fname][key]['x']
                y=self.data[fname][key]['y']
                yerr=self.data[fname][key]['yerr']
                # if len(np.where(self.data[fname][key]['yerr']<1e-30)[0])>0:
                #     QMessageBox.warning(self,'Zero Errorbars','Some or all the errorbars of the selected data are zeros.\
                #      Please select None for the Errorbar column in the Plot options of the Data_Dialog',QMessageBox.Ok)
                #     break
            # if self.fitScaleComboBox.currentText()=='Log' and len(np.where(self.data[fname]['y']<1e-30)[0])>0:
            #     posval=np.argwhere(self.fit.y>0)
            #     self.fit.y=self.data[fname]['y'][posval].T[0]
            #     self.fit.x=self.data[fname]['x'][posval].T[0]
            #     self.fit.yerr=self.data[fname]['yerr'][posval].T[0]
            self.fit.set_x(x,y=y,yerr=yerr)
            #self.update_plot()
            self.oldParams=copy.copy(self.fit.params)
            self.fit_stopped=False
            if self.fit.params['__mpar__']!={}:
                self.oldmpar=copy.deepcopy(self.mfitParamData)
            try:
                self.showFitInfoDlg(emcee_walker=emcee_walker,emcee_steps=emcee_steps, emcee_burn = emcee_burn)
                self.runFit(emcee_walker=emcee_walker, emcee_steps=emcee_steps, emcee_burn=emcee_burn,
                            emcee_cores=emcee_cores, reuse_sampler=reuse_sampler, emcee_thin=emcee_thin)
                if self.fit_stopped:
                    self.fit.result.params = self.temp_params
                #self.fit_report,self.fit_message=self.fit.perform_fit(self.xmin,self.xmax,fit_scale=self.fit_scale,\
                # fit_method=self.fit_method,callback=self.fitCallback)

                self.fit_info='Fit Message: %s\n'%self.fit_message

                self.closeFitInfoDlg()
                if self.fit_method != 'emcee':
                    self.errorAvailable=False
                    self.emcee_burn=0
                    self.emcee_steps=100
                    self.emcee_frac=self.emcee_burn/self.emcee_steps
                    self.showConfIntervalButton.setDisabled(True)
                    self.fit.functionCalled.disconnect()
                    try:
                        self.sfitParamTableWidget.cellChanged.disconnect()
                        for i in range(self.mfitParamTabWidget.count()):
                            mkey = self.mfitParamTabWidget.tabText(i)
                            self.mfitParamTableWidget[mkey].cellChanged.disconnect()
                    except:
                        pass
                    for row in range(self.sfitParamTableWidget.rowCount()):
                        key=self.sfitParamTableWidget.item(row,0).text()
                        self.sfitParamTableWidget.item(row,1).setText(self.format%(self.fit.result.params[key].value))
                        try:
                            if self.fit.result.params[key].stderr is None:
                                self.fit.result.params[key].stderr = 0.0
                            self.sfitParamTableWidget.item(row, 1).setToolTip(
                                (key + ' = ' + self.format + ' \u00B1 ' + self.format) % \
                                (self.fit.result.params[key].value,
                                 self.fit.result.params[key].stderr))
                        except:
                            pass
                    self.sfitParamTableWidget.resizeRowsToContents()
                    self.sfitParamTableWidget.resizeColumnsToContents()
                    for i in range(self.mfitParamTabWidget.count()):
                        mkey=self.mfitParamTabWidget.tabText(i)
                        for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                            for col in range(1,self.mfitParamTableWidget[mkey].columnCount()):
                                parkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                                key='__%s_%s_%03d'%(mkey,parkey,row)
                                self.mfitParamTableWidget[mkey].item(row,col).setText(self.format%(self.fit.result.params[key].value))
                                if self.fit.result.params[key].stderr is None:
                                    self.fit.result.params[key].stderr = 0.0
                                self.mfitParamTableWidget[mkey].item(row, col).setToolTip(
                                    (key + ' = ' + self.format + ' \u00B1 ' + self.format) % \
                                    (self.fit.result.params[key].value,
                                     self.fit.result.params[key].stderr))
                        self.mfitParamTableWidget[mkey].resizeRowsToContents()
                        self.mfitParamTableWidget[mkey].resizeColumnsToContents()
                    self.update_plot()
                    fitResultDlg=FitResultDialog(fit_report=self.fit_report,fit_info=self.fit_info)
                    #ans=QMessageBox.question(self,'Accept fit results?',self.fit_report,QMessageBox.Yes, QMessageBox.No)
                    if fitResultDlg.exec_():
                        for i in range(self.mfitParamTabWidget.count()):
                            mkey=self.mfitParamTabWidget.tabText(i)
                            for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                                for col in range(1, self.mfitParamTableWidget[mkey].columnCount()):
                                    parkey = self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                                    key = '__%s_%s_%03d' % (mkey, parkey, row)
                                    self.mfitParamData[mkey][parkey][row] = self.fit.result.params[key].value
                        ofname=os.path.splitext(fname.split('<>')[1])[0]
                        header='Data fitted with model: %s on %s\n'%(self.funcListWidget.currentItem().text(),time.asctime())
                        header+='Fixed Parameters\n'
                        header+='----------------\n'
                        for key in self.fit.params.keys():
                            if key not in self.fit.fit_params.keys() and key not in self.special_keys and key[:2]!='__':
                                header+=key+'='+str(self.fit.params[key])+'\n'
                        header+=self.fit_report+'\n'
                        header+="col_names=['x','y','yerr','yfit']\n"
                        header+='x \t y\t yerr \t yfit\n'
                        if type(self.fit.x)==dict:
                            for key in self.fit.x.keys():
                                fitdata=np.vstack((self.fit.x[key][self.fit.imin[key]:self.fit.imax[key]+1],
                                                   self.fit.y[key][self.fit.imin[key]:self.fit.imax[key]+1],
                                                   self.fit.yerr[key][self.fit.imin[key]:self.fit.imax[key]+1],self.fit.yfit[key])).T
                                np.savetxt(ofname+'_'+key+'_fit.txt',fitdata,header=header,comments='#')
                        else:
                            fitdata = np.vstack((self.fit.x[self.fit.imin:self.fit.imax + 1],
                                                 self.fit.y[self.fit.imin:self.fit.imax + 1],
                                                 self.fit.yerr[self.fit.imin:self.fit.imax + 1],
                                                 self.fit.yfit)).T
                            np.savetxt(ofname + '_fit.txt', fitdata, header=header, comments='#')
                        self.calcConfInterButton.setEnabled(True)
                        self.update_plot()
                        if self.autoSaveGenParamCheckBox.isChecked():
                            self.saveGenParameters(
                                bfname=os.path.join(os.path.dirname(ofname), 'genParam_' + os.path.basename(ofname)))
                        # self.xChanged()
                    else:
                        self.undoFit()
                        self.calcConfInterButton.setDisabled(True)
                    self.reuse_sampler=False
                else:
                    self.errorAvailable = True
                    self.reuse_sampler = True
                    self.emceeConfIntervalWidget.reuseSamplerCheckBox.setEnabled(True)
                    self.emceeConfIntervalWidget.reuseSamplerCheckBox.setCheckState(Qt.Checked)
                    self.fit.functionCalled.disconnect()
                    self.perform_post_sampling_tasks()
                    # self.showConfIntervalButton.setEnabled(True)
            except:
                try:
                    self.closeFitInfoDlg()
                except:
                    pass
                QMessageBox.warning(self,'Minimization failed','Some of the parameters have got unreasonable values.\n'+
                                             traceback.format_exc(),QMessageBox.Ok)
                self.update_plot()
                break
        self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
        for i in range(self.mfitParamTabWidget.count()):
            mkey=self.mfitParamTabWidget.tabText(i)
            self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
        try:
            self.fit.functionCalled.disconnect()
        except:
            pass


    def calcConfInterval(self):
        if self.confIntervalMethodComboBox.currentText()=='ChiSqrDist':
            self.confInterval_ChiSqrDist()
        else:
            self.confInterval_emcee()


    def confInterval_ChiSqrDist(self):
        self.fit_method = self.fitMethods[self.fitMethodComboBox.currentText()]
        self.confIntervalWidget=QWidget()
        self.confIntervalWidget.setWindowModality(Qt.ApplicationModal)
        uic.loadUi('./UI_Forms/ConfInterval_ChiSqrDist.ui',self.confIntervalWidget)
        self.confIntervalWidget.setWindowTitle("ChiSqrDist Confidence Interval Calculator")
        self.chidata={}
        fitTableWidget = self.confIntervalWidget.fitParamTableWidget
        self.calcErrPushButtons={}
        self.errProgressBars={}
        self.plotErrPushButtons={}
        self.stopCalc=False
        for fpar in self.fit.result.params.keys():
            if self.fit.fit_params[fpar].vary:
                row = fitTableWidget.rowCount()
                fitTableWidget.insertRow(row)
                fitTableWidget.setCellWidget(row,0,QLabel(fpar))
                fitTableWidget.setItem(row,1,QTableWidgetItem(self.format%self.fit.result.params[fpar].value))
                if self.fit.result.params[fpar].stderr is not None and self.fit.result.params[fpar].stderr!=0.0:
                    errper=5*self.fit.result.params[fpar].stderr*100/self.fit.result.params[fpar].value
                    fitTableWidget.setItem(row,2,QTableWidgetItem('%.3f' % (errper)))
                    fitTableWidget.setItem(row,3,QTableWidgetItem('%.3f' % (errper)))
                else:
                    fitTableWidget.setItem(row, 2, QTableWidgetItem('%.3f' % 10))
                    fitTableWidget.setItem(row, 3, QTableWidgetItem('%.3f' % 10))
                fitTableWidget.setItem(row,4, QTableWidgetItem('20'))
                self.calcErrPushButtons[fpar]=QPushButton('Calculate')
                fitTableWidget.setCellWidget(row, 5, self.calcErrPushButtons[fpar])
                self.calcErrPushButtons[fpar].clicked.connect(partial(self.calcErrPushButtonClicked,row,fpar))
                self.errProgressBars[fpar]=QProgressBar()
                fitTableWidget.setCellWidget(row, 6, self.errProgressBars[fpar])
                self.confIntervalWidget.fitParamTableWidget.setItem(row, 7, QTableWidgetItem(''))
                self.confIntervalWidget.fitParamTableWidget.setItem(row, 8, QTableWidgetItem(''))
                self.plotErrPushButtons[fpar]=QPushButton('Plot')
                fitTableWidget.setCellWidget(row,9, self.plotErrPushButtons[fpar])
                self.plotErrPushButtons[fpar].clicked.connect(partial(self.plotErrPushButtonClicked,row,fpar))
        fitTableWidget.resizeColumnsToContents()
        self.confIntervalWidget.plotAllPushButton.clicked.connect(self.plotAllErrPushButtonClicked)
        self.confIntervalWidget.stopPushButton.clicked.connect(self.stopErrCalc)
        self.confIntervalWidget.calcAllPushButton.clicked.connect(self.calcAllErr)
        self.confIntervalWidget.saveAllPushButton.clicked.connect(self.saveAllErr)
        self.confIntervalWidget.confIntervalSpinBox.valueChanged.connect(self.setTargetChiSqr)
        self.confIntervalWidget.saveErrPushButton.clicked.connect(self.saveParIntervalErr)
        self.minimafitparameters = copy.copy(self.fit.result.params)
        self.confIntervalWidget.showMaximized()
        self.left_limit={}
        self.right_limit={}
        self.min_value={}
        self.calcAll=False

    def stopErrCalc(self):
        self.stopCalc=True

    def setTargetChiSqr(self):
        self.confInterval = self.confIntervalWidget.confIntervalSpinBox.value()
        self.minchisqr = self.fit.result.redchi
        self.confIntervalWidget.minChiSqrLineEdit.setText(self.format % self.minchisqr)
        self.targetchisqr = self.fit.result.redchi * chi2.isf((1.0 - self.confInterval * 0.01),
                                                              self.fit.result.nfree) / (self.fit.result.nfree)
        self.confIntervalWidget.targetChiSqrLineEdit.setText(self.format % self.targetchisqr)


    def calcAllErr(self):
        self.calcAll=True
        self.stopCalc=False
        for row in range(self.confIntervalWidget.fitParamTableWidget.rowCount()):
            if not self.stopCalc:
                fpar=self.confIntervalWidget.fitParamTableWidget.cellWidget(row,0).text()
                self.calcErrPushButtonClicked(row,fpar)
            else:
                return
        self.plotAllErrPushButtonClicked()
        self.errInfoTable = []
        for key in self.chidata.keys():
            if self.left_limit[key] is not None and self.right_limit[key] is not None:
                self.errInfoTable.append([key, self.min_value[key], self.left_limit[key] - self.min_value[key],
                         self.right_limit[key] - self.min_value[key]])
            elif self.left_limit[key] is None and self.right_limit[key] is not None:
                self.errInfoTable.append([key, self.min_value[key], None,
                                          self.right_limit[key] - self.min_value[key]])
            elif self.left_limit[key] is not None and self.right_limit is None:
                self.errInfoTable.append([key, self.min_value[key], self.left_limit[key] - self.min_value[key],
                                          None])
            else:
                self.errInfoTable.append([key, self.min_value[key], None, None])


        self.confIntervalWidget.errInfoTextEdit.clear()
        self.confIntervalWidget.errInfoTextEdit.setFont(QFont("Courier", 10))
        self.confIntervalWidget.errInfoTextEdit.append(tabulate(self.errInfoTable,
                                                                headers=["Parameter","Parameter-Value","Left-Error","Right-Error"],
                                                                stralign='left',numalign='left',tablefmt='simple'))
        self.calcAll=False


    def checkMinMaxErrLimits(self,fpar,vmin,vmax):
        self.fit.fit_params[fpar].vary=False
        for key in self.minimafitparameters:  # Putting back the minima parameters
            self.fit.fit_params[key].value = self.minimafitparameters[key].value
        self.fit.fit_params[fpar].value = vmin
        fit_report, mesg = self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale,
                                                fit_method=self.fit_method,
                                                maxiter=int(self.fitIterationLineEdit.text()))
        if self.fit.result.redchi>self.targetchisqr or self.fit.fit_params[fpar].min>vmin:
            left_limit_ok=True
        else:
            left_limit_ok=False
        for key in self.minimafitparameters:  # Putting back the minima parameters
            self.fit.fit_params[key].value = self.minimafitparameters[key].value
        self.fit.fit_params[fpar].value = vmax
        fit_report, mesg = self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale,
                                                fit_method=self.fit_method,
                                                maxiter=int(self.fitIterationLineEdit.text()))
        if self.fit.result.redchi>self.targetchisqr or self.fit.fit_params[fpar].max<vmax:
            right_limit_ok=True
        else:
            right_limit_ok=False
        self.fit.fit_params[fpar].vary=True
        return left_limit_ok, right_limit_ok


    def calcErrPushButtonClicked(self,row,fpar):
        self.stopCalc=False
        for key in self.minimafitparameters:
            self.fit.fit_params[key].value = self.minimafitparameters[key].value

        self.fit.fit_params[fpar].vary=False
        redchi_r=[]
        self.errProgressBars[fpar].setMinimum(0)
        Nval = int(self.confIntervalWidget.fitParamTableWidget.item(row, 4).text())
        self.errProgressBars[fpar].setMaximum(Nval)
        #Getting the chi-sqr value at the minima position keeping the value of fpar fixed at the minima position
        fit_report, mesg =self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale, fit_method=self.fit_method,
                             maxiter=int(self.fitIterationLineEdit.text()))
        self.setTargetChiSqr()
        redchi_r.append([self.fit.fit_params[fpar].value, self.fit.result.redchi])
        self.errProgressBars[fpar].setValue(1)
        value=self.fit.result.params[fpar].value
        vmax = value*(1.0+float(self.confIntervalWidget.fitParamTableWidget.item(row, 3).text())/100.0)
        vmin = value*(1.0-float(self.confIntervalWidget.fitParamTableWidget.item(row, 2).text())/100.0)

        left_limit_ok,right_limit_ok=self.checkMinMaxErrLimits(fpar,vmin,vmax)
        self.fit.fit_params[fpar].vary = False
        if left_limit_ok and right_limit_ok:
            # Fitting the right hand side of the minima starting from the first point after minima
            self.min_value[fpar]=value
            pvalues=np.linspace(value+(vmax-value)*2/Nval, vmax, int(Nval/2))
            i=1
            for parvalue in pvalues:
                if self.stopCalc:
                    for key in self.minimafitparameters:
                        self.fit.fit_params[key].value = self.minimafitparameters[key].value
                    return
                for key in self.minimafitparameters: # Putting back the minima parameters
                    self.fit.fit_params[key].value = self.minimafitparameters[key].value
                self.fit.fit_params[fpar].value=parvalue
                fit_report, mesg = self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale,
                                                        fit_method=self.fit_method,
                                                        maxiter=int(self.fitIterationLineEdit.text()))
                if self.fit.result.success:
                    redchi_r.append([parvalue,self.fit.result.redchi])
                i+=1
                self.errProgressBars[fpar].setValue(i)
                QApplication.processEvents()


            step=(value-vmin)*2/Nval
            redchi_l=[redchi_r[0]]

            #Fitting the left hand of the minima starting from the minima point
            pvalues=np.linspace(value-step, vmin, int(Nval / 2))
            for parvalue in pvalues:
                if self.stopCalc:
                    for key in self.minimafitparameters:
                        self.fit.fit_params[key].value = self.minimafitparameters[key].value
                    return
                for key in self.minimafitparameters: # Putting back the minima parameters
                    self.fit.fit_params[key].value = self.minimafitparameters[key].value
                self.fit.fit_params[fpar].value = parvalue
                fit_report, mesg = self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale,
                                                        fit_method=self.fit_method,
                                                        maxiter=int(self.fitIterationLineEdit.text()))
                if self.fit.result.success:
                    redchi_l.append([parvalue, self.fit.result.redchi])
                i+=1
                self.errProgressBars[fpar].setValue(i)
                QApplication.processEvents()

            chidata=np.array(redchi_r+redchi_l[1:])
            self.chidata[fpar]=chidata[chidata[:,0].argsort()]


            # Calculating the right-limit by interpolation
            rvalues = np.array(redchi_r)
            if self.targetchisqr < np.max(rvalues[:, 1]):
                fn=interp1d(rvalues[:, 1], rvalues[:, 0],kind='linear')
                self.right_limit[fpar] = fn(self.targetchisqr)
                self.confIntervalWidget.fitParamTableWidget.item(row, 8).setText(self.format % (self.right_limit[fpar]))
            else:
                self.right_limit[fpar] = None
                self.confIntervalWidget.fitParamTableWidget.item(row, 8).setText('None')

            # Calculating the left-limit by interpolation
            lvalues = np.array(redchi_l)
            if self.targetchisqr < np.max(lvalues[:, 1]):
                fn=interp1d(lvalues[:, 1], lvalues[:, 0],kind='linear')
                self.left_limit[fpar] = fn(self.targetchisqr)
                self.confIntervalWidget.fitParamTableWidget.item(row, 7).setText(self.format % (self.left_limit[fpar]))
            else:
                self.left_limit[fpar] = None
                self.confIntervalWidget.fitParamTableWidget.item(row, 7).setText('None')
            self.confIntervalWidget.fitParamTableWidget.resizeColumnsToContents()

            # Plotting the data
            if not self.calcAll:
                self.plotErrPushButtonClicked(row, fpar)

            #Showing the Errorbars
            self.errInfoTable = []
            key=fpar
            if self.left_limit[key] is not None and self.right_limit[key] is not None:
                self.errInfoTable.append([key, self.min_value[key], self.left_limit[key] - self.min_value[key],
                                          self.right_limit[key] - self.min_value[key]])
            elif self.left_limit[key] is None and self.right_limit[key] is not None:
                self.errInfoTable.append([key, self.min_value[key], None,
                                          self.right_limit[key] - self.min_value[key]])
            elif self.left_limit[key] is not None and self.right_limit is None:
                self.errInfoTable.append([key, self.min_value[key], self.left_limit[key] - self.min_value[key],
                                          None])
            else:
                self.errInfoTable.append([key, self.min_value[key], None, None])
            self.confIntervalWidget.errInfoTextEdit.clear()
            self.confIntervalWidget.errInfoTextEdit.setFont(QFont("Courier", 10))
            self.confIntervalWidget.errInfoTextEdit.append(tabulate(self.errInfoTable,
                                                                    headers=["Parameter", "Parameter-Value",
                                                                             "Left-Error", "Right-Error"],
                                                                    stralign='left', numalign='left',
                                                                    tablefmt='simple'))

        elif left_limit_ok:
            QMessageBox.warning(self,'Limit Warning','Max limit is not enough to reach the target chi-square for %s. Increase the Max limit'%fpar,QMessageBox.Ok)
            self.errProgressBars[fpar].setValue(0)
            QApplication.processEvents()
        else:
            QMessageBox.warning(self, 'Limit Warning', 'Min limit is not enough to reach the target chi-square for %s. Increase the Min limit'%fpar, QMessageBox.Ok)
            self.errProgressBars[fpar].setValue(0)
            QApplication.processEvents()

        # Going back to the minimum chi-sqr condition
        for key in self.minimafitparameters:
            self.fit.fit_params[key].value = self.minimafitparameters[key].value
        self.fit.fit_params[fpar].vary = True
        fit_report, mesg = self.fit.perform_fit(self.xmin, self.xmax, fit_scale=self.fit_scale,
                                                fit_method=self.fit_method,
                                                maxiter=int(self.fitIterationLineEdit.text()))

    def plotErrPushButtonClicked(self,row,fpar):
        if fpar in self.chidata.keys():
            mw=MplWidget()
            mw.setWindowModality(Qt.ApplicationModal)
            subplot=mw.getFigure().add_subplot(111)
            subplot.plot(self.chidata[fpar][:, 0], self.chidata[fpar][:, 1], 'r.')
            subplot.axhline(self.minchisqr,color='k',lw=1,ls='--')
            subplot.axhline(self.targetchisqr,color='k',lw=1,ls='-')
            subplot.axvline(self.min_value[fpar],color='b',lw=2,ls='-')
            # pl.text(self.min_value[fpar],1.01*self.minchisqr,self.format%self.min_value[fpar],rotation='vertical')
            if self.right_limit[fpar] is not None:
                subplot.axvline(self.right_limit[fpar],color='b',lw=1,ls='--')
                # pl.text(self.right_limit[fpar], 1.01*self.targetchisqr, self.format%self.right_limit[fpar],rotation='vertical')
                right_error = self.right_limit[fpar]-self.min_value[fpar]
            else:
                right_error='None'
            if self.left_limit[fpar] is not None:
                subplot.axvline(self.left_limit[fpar],color='b',lw=1,ls='--')
                # pl.text(self.left_limit[fpar], 1.01*self.targetchisqr, self.format% self.left_limit[fpar],rotation='vertical')
                left_error = self.left_limit[fpar]-self.min_value[fpar]
            else:
                left_error='None'
            subplot.set_title('%.3e$^{%.3e}_{%.3e}$'%(self.min_value[fpar], right_error, left_error))
            subplot.set_xlabel(fpar)
            subplot.set_ylabel('\u03c7$^2$')
            mw.getFigure().tight_layout()
            mw.draw()
            mw.show()
        else:
            QMessageBox.warning(self, 'Data error', 'No data available for plotting. Calculate first', QMessageBox.Ok)


    def plotAllErrPushButtonClicked(self):
        pkey=list(self.chidata.keys())
        Nplots=len(pkey)
        if Nplots>0:
            mw=MplWidget()
            mw.setWindowModality(Qt.ApplicationModal)
            rows=math.ceil(np.sqrt(Nplots))
            i=1
            for row in range(rows):
                for col in range(rows):
                    if i<=Nplots:
                        ax=mw.getFigure().add_subplot(rows,rows,i)
                        ax.plot(self.chidata[pkey[i-1]][:,0],self.chidata[pkey[i-1]][:,1],'r.')
                        ax.axhline(self.minchisqr, color='k', lw=1, ls='--')
                        ax.axhline(self.targetchisqr, color='k', lw=1, ls='-')
                        ax.axvline(self.min_value[pkey[i-1]], color='b', lw=2, ls='-')
                        # ax[row,col].text(self.min_value[pkey[i-1]], 1.01 * self.minchisqr, self.format % self.min_value[pkey[i-1]],rotation='vertical')
                        if self.right_limit[pkey[i-1]] is not None:
                            ax.axvline(self.right_limit[pkey[i-1]], color='b', lw=1, ls='--')
                            right_error=self.right_limit[pkey[i-1]]-self.min_value[pkey[i-1]]
                            # ax[row,col].text(self.right_limit[pkey[i-1]], 1.01*self.targetchisqr, self.format % self.right_limit[pkey[i-1]],rotation='vertical')
                        else:
                            right_error='None'
                        if self.left_limit[pkey[i-1]] is not None:
                            ax.axvline(self.left_limit[pkey[i-1]], color='b', lw=1, ls='--')
                            left_error=self.left_limit[pkey[i-1]]-self.min_value[pkey[i-1]]
                            # ax[row, col].text(self.left_limit[pkey[i-1]], 1.01*self.targetchisqr, self.format % self.left_limit[pkey[i-1]],rotation='vertical')
                        else:
                            left_error='None'
                        ax.set_title('%.3e$^{%.3e}_{%.3e}$'%(self.min_value[pkey[i-1]], right_error,left_error))
                        ax.set_xlabel(pkey[i-1])
                        ax.set_ylabel('\u03c7$^2$')
                    i+=1
            mw.getFigure().tight_layout()
            mw.draw()
            mw.show()

    def saveAllErr(self):
        fname=QFileDialog.getSaveFileName(self,'Provide prefix of the filename',directory=self.curDir,filter='Chi-Sqr files (*.chisqr)')[0]
        if fname!='':
            for key in self.chidata.keys():
                filename=os.path.splitext(fname)[0]+'_'+key+'.chisqr'
                header='Saved on %s\n'%(time.asctime())
                header="col_names=['%s','chi-sqr']\n"%key
                header+='%s\tchi-sqr'%key
                pl.savetxt(filename,self.chidata[key],header=header)

    def saveParIntervalErr(self):
        fname = QFileDialog.getSaveFileName(caption='Save Parameter Errors as', filter='Parameter Error files (*.perr)',
                                            directory=self.curDir)[0]
        if fname!='':
            fh=open(fname,'w')
            fh.write('# File saved on %s\n'%time.asctime())
            fh.write('# Error calculated using Chi-Sqr-Distribution Method\n')
            tlines=tabulate(self.errInfoTable, headers=["Parameter","Parameter-Value","Left-Error","Right-Error"],
                                   stralign='left',numalign='left',tablefmt='simple')
            lines=tlines.split('\n')
            for i,line in enumerate(lines):
                if i<2:
                    fh.write('#'+line+'\n')
                else:
                    fh.write(' '+line+'\n')
            fh.close()




        
    def confInterval_emcee(self):
        """
        """
        self.fit_method = self.fitMethods[self.fitMethodComboBox.currentText()]
        if not self.errorAvailable:
             self.emcee_walker=(self.fit.result.nvarys+1)*2
        else:
        #     # try:
             tnum=len(self.fit.result.flatchain[self.fit.result.var_names[0]])/self.emcee_walker
             self.emcee_frac=self.emcee_burn/(tnum/(1.0-self.emcee_frac))
             emcee_burn=tnum*self.emcee_frac/(1.0-self.emcee_frac)
             self.emcee_burn=int(emcee_burn+self.emcee_steps*self.emcee_frac)
        self.emceeConfIntervalWidget = QWidget()
        self.emceeConfIntervalWidget.setWindowModality(Qt.ApplicationModal)
        uic.loadUi('./UI_Forms/EMCEE_ConfInterval_Widget.ui', self.emceeConfIntervalWidget)
        self.emceeConfIntervalWidget.setWindowTitle('MCMC Confidence Interval Caclulator')
        self.emceeConfIntervalWidget.MCMCWalkerLineEdit.setText(str(self.emcee_walker))
        self.emceeConfIntervalWidget.MCMCStepsLineEdit.setText(str(self.emcee_steps))
        self.emceeConfIntervalWidget.MCMCBurnLineEdit.setText(str(self.emcee_burn))
        self.emceeConfIntervalWidget.MCMCThinLineEdit.setText(str(self.emcee_thin))
        self.emceeConfIntervalWidget.ParallelCoresLineEdit.setText(str(self.emcee_cores))
        if not self.errorAvailable:
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setChecked(False)
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setDisabled(True)
            self.reuse_sampler=False
        else:
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setChecked(True)
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setDisabled(True)

        if self.reuse_sampler:
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setEnabled(True)
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setCheckState(Qt.Checked)
        else:
            self.emceeConfIntervalWidget.reuseSamplerCheckBox.setCheckState(Qt.Unchecked)

        self.emceeConfIntervalWidget.startSamplingPushButton.clicked.connect(self.start_emcee_sampling)
        self.emceeConfIntervalWidget.MCMCWalkerLineEdit.returnPressed.connect(self.MCMCWalker_changed)
        self.emceeConfIntervalWidget.saveConfIntervalPushButton.clicked.connect(self.saveParameterError)
        self.emceeConfIntervalWidget.addUserDefinedParamPushButton.clicked.connect(lambda x:
                                                                                   self.addMCMCUserDefinedParam(parname=None, expression=None))
        self.emceeConfIntervalWidget.removeUserDefinedParamPushButton.clicked.connect(self.removeMCMCUserDefinedParam)
        self.emceeConfIntervalWidget.saveUserDefinedParamPushButton.clicked.connect(self.saveMCMCUserDefinedParam)
        self.emceeConfIntervalWidget.loadUserDefinedParamPushButton.clicked.connect(self.loadMCMCUserDefinedParam)
        self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemDoubleClicked.connect(self.openMCMCUserDefinedParam)
        self.emceeConfIntervalWidget.progressBar.setValue(0)
        self.EnableUserDefinedParameterButtons(enable=False)
        self.emceeConfIntervalWidget.showMaximized()
        if self.errorAvailable:
            self.update_emcee_parameters()
            self.perform_post_sampling_tasks()
            self.cornerPlot()
            self.emceeConfIntervalWidget.tabWidget.setCurrentIndex=(4)

    def EnableUserDefinedParameterButtons(self,enable=False):
        self.emceeConfIntervalWidget.addUserDefinedParamPushButton.setEnabled(enable)
        self.emceeConfIntervalWidget.removeUserDefinedParamPushButton.setEnabled(enable)
        self.emceeConfIntervalWidget.saveUserDefinedParamPushButton.setEnabled(enable)
        self.emceeConfIntervalWidget.loadUserDefinedParamPushButton.setEnabled(enable)



    def openMCMCUserDefinedParam(self,item,column):
        txt=item.text(0)
        if ':chain:' not in txt:
            parname,expression=txt.split('=')
            self.addMCMCUserDefinedParam(parname=parname, expression=expression)


    def addMCMCUserDefinedParam(self, parname=None, expression=None):
        self.MCMCUserDefinedParamWidget = QWidget()
        self.MCMCUserDefinedParamWidget.setWindowModality(Qt.ApplicationModal)
        uic.loadUi('./UI_Forms/User_Defined_Param_Widget.ui', self.MCMCUserDefinedParamWidget)
        if parname is not None:
            self.MCMCUserDefinedParamWidget.parnameLineEdit.setText(parname)
            self.MCMCUserDefinedParamWidget.parnameLineEdit.setEnabled(False)
            new=False
        else:
            new=True
        if expression is not None:
            self.MCMCUserDefinedParamWidget.expressionTextEdit.setText(expression)
        availParams = self.fit.result.var_names
        self.MCMCUserDefinedParamWidget.availParamListWidget.addItems(availParams)

        self.MCMCUserDefinedParamWidget.addPushButton.clicked.connect(lambda x: self.acceptUserDefinedParams(new=new))
        self.MCMCUserDefinedParamWidget.cancelPushButton.clicked.connect(self.MCMCUserDefinedParamWidget.close)
        self.MCMCUserDefinedParamWidget.availParamListWidget.itemDoubleClicked.connect(self.appendUserDefinedExpression)
        self.MCMCUserDefinedParamWidget.show()

    def appendUserDefinedExpression(self,item):
        txt=item.text()
        self.MCMCUserDefinedParamWidget.expressionTextEdit.insertPlainText(txt)
        self.MCMCUserDefinedParamWidget.expressionTextEdit.moveCursor(QTextCursor.End)

    def acceptUserDefinedParams(self,new=True,parname=None,expression=None):
        if parname is None:
            parname=self.MCMCUserDefinedParamWidget.parnameLineEdit.text()
        if parname=='':
            QMessageBox.warning(self, 'Name Error', 'Please provide a parameter name.',QMessageBox.Ok)
            return
        elif parname in self.param_chain.keys() and new:
            QMessageBox.warning(self, 'Name Error', 'Please provide a parameter name which is not already used.', QMessageBox.Ok)
            return
        else:
            if expression is None:
                txt=self.MCMCUserDefinedParamWidget.expressionTextEdit.toPlainText()
            else:
                txt=expression
            if txt != 'None':
                try:
                    self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.disconnect(
                        self.userDefinedParameterTreeSelectionChanged)
                except:
                    pass
                if new:
                    l1 = QTreeWidgetItem([parname+'='+txt])
                else:
                    self.emceeConfIntervalWidget.userDefinedParamTreeWidget.currentItem().setText(0,parname+'='+txt)
                self.param_chain[parname] = {}
                for i in range(self.chain_shape[1]):
                    ttxt=txt[0:]
                    for name in self.fit.result.var_names:
                        ttxt=ttxt.replace(name,"self.param_chain['%s'][%d]"%(name,i))
                    try:
                        self.param_chain[parname][i]=eval(ttxt)
                        if new:
                            l1_child = QTreeWidgetItem(['%s:chain:%d' % (parname, i)])
                            l1.addChild(l1_child)
                            self.emceeConfIntervalWidget.userDefinedParamTreeWidget.addTopLevelItem(l1)
                        if expression is None:
                            self.MCMCUserDefinedParamWidget.close()

                    except:
                        QMessageBox.warning(self, 'Expression Error',
                                            'Some problems in the expression\n' + traceback.format_exc(),
                                            QMessageBox.Ok)
                        return
                self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.connect(
                    self.userDefinedParameterTreeSelectionChanged)

            else:
                self.MCMCUserDefinedParamWidget.close()

    def removeMCMCUserDefinedParam(self):
        try:
            self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.disconnect()
        except:
            pass
        for item in self.emceeConfIntervalWidget.userDefinedParamTreeWidget.selectedItems():
            parname=item.text(0).split('=')[0]
            del self.param_chain[parname]
            index=self.emceeConfIntervalWidget.userDefinedParamTreeWidget.indexOfTopLevelItem(item)
            self.emceeConfIntervalWidget.userDefinedParamTreeWidget.takeTopLevelItem(index)
        self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.connect(self.userDefinedParameterTreeSelectionChanged)


    def saveMCMCUserDefinedParam(self):
        fname=QFileDialog.getSaveFileName(caption='Save User-Defined parameter expression as',filter='Expression files (*.expr)'
                                          ,directory=self.curDir)[0]
        if fname!='':
            if os.path.splitext(fname)[0]=='':
                fname=fname+'.expr'
        else:
            return
        fh=open(fname,'w')
        txt ='#Parameter Expressions saved on %s\n'%(time.asctime())
        txt += '#Category:%s\n' % self.curr_category
        txt += '#Function:%s\n' % self.funcListWidget.currentItem().text()
        root = self.emceeConfIntervalWidget.userDefinedParamTreeWidget.invisibleRootItem()
        child_count=root.childCount()
        for i in range(child_count):
            txt += '%s\n'%root.child(i).text(0)
        fh.write(txt)
        fh.close()


    def loadMCMCUserDefinedParam(self):
        fname=QFileDialog.getOpenFileName(self,'Open Expression file',filter='Expression files (*.expr)'
                                          ,directory=self.curDir)[0]
        if fname!='':
            try:
                self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.disconnect()
            except:
                pass
            #Removing the existing User-Defined parameters, if present
            root = self.emceeConfIntervalWidget.userDefinedParamTreeWidget.invisibleRootItem()
            child_count = root.childCount()
            for i in range(child_count):
                parname, expression = root.child(i).text(0).split('=')
                del self.param_chain[parname]
            self.emceeConfIntervalWidget.userDefinedParamTreeWidget.clear()

            fh=open(fname,'r')
            lines=fh.readlines()
            for line in lines:
                if line[0]=='#':
                    if 'Category' in line:
                        _,category=line[1:].strip().split(':')
                        if category!=self.curr_category:
                            QMessageBox.warning(self,'File Error','The expression file is not generated from the same category of functions as used here.'
                                                ,QMessageBox.Ok)
                            return
                    elif 'Function' in line:
                        _,func=line[1:].strip().split(':')
                        if func!=self.curr_module:
                            QMessageBox.warning(self, 'File Error',
                                                'The expression file is not generated from the same function as used here.'
                                                , QMessageBox.Ok)
                            return
                else:
                    parname,expression=line.strip().split('=')
                    self.acceptUserDefinedParams(new=True,parname=parname,expression=expression)
        self.emceeConfIntervalWidget.userDefinedParamTreeWidget.itemSelectionChanged.connect(self.userDefinedParameterTreeSelectionChanged)




    def MCMCWalker_changed(self):
        self.emceeConfIntervalWidget.reuseSamplerCheckBox.setCheckState(Qt.Unchecked)
        self.update_emcee_parameters()


    def update_emcee_parameters(self):
        self.emcee_walker=int(self.emceeConfIntervalWidget.MCMCWalkerLineEdit.text())
        self.emcee_steps=int(self.emceeConfIntervalWidget.MCMCStepsLineEdit.text())
        self.emcee_burn=int(self.emceeConfIntervalWidget.MCMCBurnLineEdit.text())
        self.emcee_thin = int(self.emceeConfIntervalWidget.MCMCThinLineEdit.text())
        if self.emceeConfIntervalWidget.reuseSamplerCheckBox.isChecked():
            self.reuse_sampler=True
        else:
            self.reuse_sampler=False
        self.emcee_cores = int(self.emceeConfIntervalWidget.ParallelCoresLineEdit.text())

    def start_emcee_sampling(self):
        try:
            self.emceeConfIntervalWidget.parameterTreeWidget.itemSelectionChanged.disconnect()
        except:
            pass
        self.emceeConfIntervalWidget.parameterTreeWidget.clear()
        self.emceeConfIntervalWidget.chainMPLWidget.clear()
        self.emceeConfIntervalWidget.correlationMPLWidget.clear()
        self.emceeConfIntervalWidget.cornerPlotMPLWidget.clear()
        self.emceeConfIntervalWidget.confIntervalTextEdit.clear()
        self.update_emcee_parameters()
        if not self.errorAvailable:
            self.emcee_frac=self.emcee_burn/self.emcee_steps
        self.doFit(fit_method='emcee', emcee_walker=self.emcee_walker, emcee_steps=self.emcee_steps,
                       emcee_cores=self.emcee_cores, reuse_sampler=self.reuse_sampler, emcee_burn=self.emcee_burn,
                   emcee_thin=self.emcee_thin)


    def conf_interv_status(self,params,iterations,residual,fit_scale):
        self.confIntervalStatus.setText(self.confIntervalStatus.text().split('\n')[0]+'\n\n {:^s} = {:10d}'.format('Iteration',iterations))            
        QApplication.processEvents()
        
    def runFit(self,  emcee_walker=100, emcee_steps=100, emcee_cores=1, reuse_sampler=False, emcee_burn=30, emcee_thin=1):
        self.start_time=time.time()
        self.fit_report,self.fit_message=self.fit.perform_fit(self.xmin,self.xmax,fit_scale=self.fit_scale, fit_method=self.fit_method,
                                                              maxiter=int(self.fitIterationLineEdit.text()),
                                                              emcee_walker=emcee_walker, emcee_steps=emcee_steps,
                                                              emcee_cores=emcee_cores, reuse_sampler=reuse_sampler, emcee_burn=emcee_burn,
                                                              emcee_thin=emcee_thin)
        
    
    def showFitInfoDlg(self, emcee_walker=100, emcee_steps=100, emcee_burn=30):
        if self.fit_method!='emcee':
            self.fitInfoDlg=QDialog(self)
            vblayout=QVBoxLayout(self.fitInfoDlg)
            self.fitIterLabel=QLabel('Iteration: 0,\t Chi-Sqr: Not Available',self.fitInfoDlg)
            vblayout.addWidget(self.fitIterLabel)
            self.stopFitPushButton=QPushButton('Stop')
            vblayout.addWidget(self.stopFitPushButton)
            self.stopFitPushButton.clicked.connect(self.stopFit)
            self.fitInfoDlg.setWindowTitle('Please wait for the fitting to be completed')
            self.fitInfoDlg.setModal(True)
            self.fitInfoDlg.show()
        else:
            self.emceeConfIntervalWidget.fitIterLabel.setText('Time left (hh:mm:ss): %s'%('N.A.'))
            self.emceeConfIntervalWidget.progressBar.setMaximum(emcee_walker*emcee_steps)
            self.emceeConfIntervalWidget.progressBar.setMinimum(0)
            self.emceeConfIntervalWidget.progressBar.setValue(0)
            self.emceeConfIntervalWidget.stopSamplingPushButton.clicked.connect(self.stopFit)

    def stopFit(self):
        self.fit.fit_abort=True
        self.fit_stopped=True
        self.reuse_sampler=False
        if self.fit_method=='emcee':
            self.emceeConfIntervalWidget.stopSamplingPushButton.clicked.disconnect()

    def closeFitInfoDlg(self):
        self.fitInfoDlg.done(0)

        

    def fitCallback(self,params,iterations,residual,fit_scale):
        # self.fitIterLabel.setText('Iteration=%d,\t Chi-Sqr=%.5e'%(iterations,np.sum(residual**2)))
        # if np.any(self.fit.yfit):
        chisqr=np.sum(residual**2)
        if chisqr<self.tchisqr:
            self.fitIterLabel.setText('Iteration=%d,\t Chi-Sqr=%.5e' % (iterations,chisqr))
            self.temp_params=copy.copy(params)
            if type(self.fit.x)==dict:
                for key in self.fit.x.keys():
                    self.plotWidget.add_data(x=self.fit.x[key][self.fit.imin[key]:self.fit.imax[key]+1],y=self.fit.yfit[key],\
                                     name=self.funcListWidget.currentItem().text()+':'+key,fit=True)
                    self.fit.params['output_params']['Residuals_%s'%key] = {'x': self.fit.x[key][self.fit.imin[key]:self.fit.imax[key]+1],
                                                                            'y': (self.fit.y[key][self.fit.imin[key]:self.fit.imax[key]+1]-self.fit.yfit[key])
                    /self.fit.yerr[key][self.fit.imin[key]:self.fit.imax[key]+1]}
            else:
                self.plotWidget.add_data(x=self.fit.x[self.fit.imin:self.fit.imax + 1], y=self.fit.yfit, \
                                         name=self.funcListWidget.currentItem().text(), fit=True)
            # else:
            #     QMessageBox.warning(self,'Parameter Value Error','One or more fitting parameters has got unphysical values perhaps to make all the yvalues zeros!',QMessageBox.Ok)
            #     self.fit.fit_abort=True
                self.fit.params['output_params']['Residuals']={'x':self.fit.x[self.fit.imin:self.fit.imax + 1],
                                                               'y': (self.fit.y[self.fit.imin:self.fit.imax + 1]-self.fit.yfit)/self.fit.yerr[self.fit.imin:self.fit.imax + 1]}
            self.tchisqr=chisqr
        QApplication.processEvents()


    def fitErrorCallback(self, params, iterations, residual, fit_scale):
        time_taken=time.time()-self.start_time
        frac=iterations/(self.emcee_walker*self.emcee_steps+self.emcee_walker)
        time_left=time_taken*(self.emcee_walker*self.emcee_steps+self.emcee_walker-iterations)/iterations
        self.emceeConfIntervalWidget.fitIterLabel.setText('Time left (hh:mm:ss): %s'%(time.strftime('%H:%M:%S',time.gmtime(time_left))))
        self.emceeConfIntervalWidget.progressBar.setValue(iterations)
        QApplication.processEvents()

    def perform_post_sampling_tasks(self):
        self.emceeConfIntervalWidget.progressBar.setValue(self.emcee_walker*self.emcee_steps)
        self.emceeConfIntervalWidget.fitIterLabel.setText('Time left (hh:mm:ss): 00:00:00' )
        self.chain=self.fit.result.chain
        self.chain_shape=self.chain.shape
        self.param_chain=OrderedDict()
        for i,key in enumerate(self.fit.result.flatchain.keys()):
            l1=QTreeWidgetItem([key])
            self.param_chain[key]=OrderedDict()
            for j in range(self.chain_shape[1]):
                self.param_chain[key][j]=self.chain[:,j,i]
                l1_child=QTreeWidgetItem(['%s:chain:%d'%(key,j)])
                l1.addChild(l1_child)
            self.emceeConfIntervalWidget.parameterTreeWidget.addTopLevelItem(l1)
        self.emceeConfIntervalWidget.parameterTreeWidget.itemSelectionChanged.connect(self.parameterTreeSelectionChanged)

        #Calculating autocorrelation
        acor=OrderedDict()
        Nrows=len(self.param_chain.keys())
        self.emceeConfIntervalWidget.correlationMPLWidget.clear()
        ax1 = self.emceeConfIntervalWidget.correlationMPLWidget.fig.add_subplot(1, 1, 1)
        corr_time=[]
        acor_mcmc=self.fit.fitter.sampler.get_autocorr_time(quiet=True)
        for i,key in enumerate(self.param_chain.keys()):
            tcor=[]
            for ikey in self.param_chain[key].keys():
                tdata=self.param_chain[key][ikey]
                res=sm.tsa.acf(tdata,nlags=len(tdata),fft=True)
                tcor.append(res)
            tcor=np.array(tcor)
            acor[key]=np.mean(tcor,axis=0)
            ax1.plot(acor[key],'-',label='para=%s'%key)
            corr_time.append([key,acor_mcmc[i]])#np.sum(np.where(acor[key]>0,acor[key],0))])
        ax1.set_xlabel('Steps')
        ax1.set_ylabel('Autocorrelation')
        l=ax1.legend(loc='best')
        l.set_draggable(True)
        self.emceeConfIntervalWidget.correlationMPLWidget.draw()
        self.emceeConfIntervalWidget.corrTimeTextEdit.clear()
        self.emceeConfIntervalWidget.corrTimeTextEdit.setFont(QFont("Courier", 10))
        corr_text = tabulate(corr_time, headers=['Parameter', 'Correlation-time (Steps)'], stralign='left',
                             numalign='left', tablefmt='simple')
        self.emceeConfIntervalWidget.corrTimeTextEdit.append(corr_text)

        #Plotting Acceptance Ratio
        self.emceeConfIntervalWidget.acceptFracMPLWidget.clear()
        ax2=self.emceeConfIntervalWidget.acceptFracMPLWidget.fig.add_subplot(1,1,1)
        ax2.plot(self.fit.result.acceptance_fraction,'-')
        ax2.set_xlabel('Walkers')
        ax2.set_ylabel('Acceptance Ratio')
        self.emceeConfIntervalWidget.acceptFracMPLWidget.draw()

        self.emceeConfIntervalWidget.calcConfIntervPushButton.clicked.connect(self.cornerPlot)
        self.emceeConfIntervalWidget.tabWidget.setCurrentIndex(2)
        self.reset_cornerplot=True

        #Calculating User-Defined parameters
        self.EnableUserDefinedParameterButtons(enable=True)
        root = self.emceeConfIntervalWidget.userDefinedParamTreeWidget.invisibleRootItem()
        child_count = root.childCount()
        if child_count>0:
            for i in range(child_count):
                parname, expression = root.child(i).text(0).split('=')
                self.param_chain[parname]={}
                for j in range(self.chain_shape[1]):
                    ttxt = expression[0:]
                    for name in self.fit.result.var_names:
                        ttxt = ttxt.replace(name, "self.param_chain['%s'][%d]" % (name, j))
                    self.param_chain[parname][j] = eval(ttxt)

    def cornerPlot(self):
        percentile = self.emceeConfIntervalWidget.percentileDoubleSpinBox.value()
        first = int(self.emceeConfIntervalWidget.MCMCBurnLineEdit.text())
        if self.reset_cornerplot:
            self.emceeConfIntervalWidget.cornerPlotMPLWidget.clear()
            names = self.fit.result.var_names#[name for name in self.fit.result.var_names if name != '__lnsigma']
            values = [self.fit.result.params[name].value for name in names]
            ndim = len(names)
            quantiles=[1.0-percentile/100,0.5,percentile/100]
            corner.corner(self.fit.result.flatchain[names][first:], labels=names, bins=50, levels=(percentile/100,),
                          truths=values, quantiles=quantiles, show_titles=True, title_fmt='.6f',
                          use_math_text=True, title_kwargs={'fontsize': 3 * 12 / ndim},
                          label_kwargs={'fontsize': 3 * 12 / ndim}, fig=self.emceeConfIntervalWidget.cornerPlotMPLWidget.fig)
            for ax3 in self.emceeConfIntervalWidget.cornerPlotMPLWidget.fig.get_axes():
                ax3.set_xlabel('')
                ax3.set_ylabel('')
                ax3.tick_params(axis='y', labelsize=3 * 12 / ndim, rotation=0)
                ax3.tick_params(axis='x', labelsize=3 * 12 / ndim)
            self.emceeConfIntervalWidget.cornerPlotMPLWidget.draw()
            self.emceeConfIntervalWidget.tabWidget.setCurrentIndex(4)
            self.reset_cornerplot=False
        self.calcMCMCerrorbars(burn=first,percentile=percentile)
        # err_quantiles={}
        # mesg = [['Parameters', 'Value(50%)', 'Left-error(%.3f)'%(100-percentile), 'Right-error(%.3f)'%percentile]]
        # for name in names:
        #     err_quantiles[name] = corner.quantile(self.fit.result.flatchain[name], quantiles)
        #     l,p,r=err_quantiles[name]
        #     mesg.append([name, p, l - p, r - p])
        #
        # self.emceeConfIntervalWidget.confIntervalTextEdit.clear()
        # self.emceeConfIntervalWidget.confIntervalTextEdit.setFont(QFont("Courier", 10))
        # txt = tabulate(mesg, headers='firstrow', stralign='left', numalign='left', tablefmt='simple')
        # self.emceeConfIntervalWidget.confIntervalTextEdit.append(txt)


    def calcMCMCerrorbars(self,burn=0,percentile=85):
        mesg = [['Parameters', 'Value(50%)', 'Left-error(%.3f%s)' % (100 - percentile,'%'), 'Right-error(%.3f%s)' % (percentile,'%')]]
        for key in self.param_chain.keys():
            for chain in self.param_chain[key].keys():
                try:
                    pardata=np.vstack((pardata,self.param_chain[key][chain][burn:]))
                except:
                    pardata=[self.param_chain[key][chain][burn:]]
            pardata=np.ndarray.flatten(pardata)
            errors=np.quantile(pardata,[(100-percentile)/100,50/100,percentile/100])
            mesg.append([key, errors[1], errors[1]-errors[0], errors[2]-errors[1]])
        self.emceeConfIntervalWidget.confIntervalTextEdit.clear()
        self.emceeConfIntervalWidget.confIntervalTextEdit.setFont(QFont("Courier", 10))
        txt = tabulate(mesg, headers='firstrow', stralign='left', numalign='left', tablefmt='simple')
        self.emceeConfIntervalWidget.confIntervalTextEdit.append(txt)


    def parameterTreeSelectionChanged(self):
        self.emceeConfIntervalWidget.chainMPLWidget.clear()
        chaindata={}
        for item in self.emceeConfIntervalWidget.parameterTreeWidget.selectedItems():
            if ':chain:' in item.text(0):
                key,i=item.text(0).split(':chain:')
                try:
                    chaindata[key].append(int(i))
                except:
                    chaindata[key]=[int(i)]
        NRows = len(chaindata.keys())
        if NRows>0:
            ax={}
            firstkey=list(chaindata.keys())[0]
            for j,key in enumerate(chaindata.keys()):
                try:
                    ax[key]=self.emceeConfIntervalWidget.chainMPLWidget.fig.add_subplot(NRows, 1, j+1, sharex=ax[firstkey])
                except:
                    ax[key] = self.emceeConfIntervalWidget.chainMPLWidget.fig.add_subplot(NRows, 1, j+1)
                for i in chaindata[key]:
                    ax[key].plot(self.param_chain[key][i],'-')
                ax[key].set_xlabel('MC steps')
                ax[key].set_ylabel(key)
            self.emceeConfIntervalWidget.chainMPLWidget.draw()
            self.emceeConfIntervalWidget.tabWidget.setCurrentIndex(0)

    def userDefinedParameterTreeSelectionChanged(self):
        self.emceeConfIntervalWidget.userDefinedChainMPLWidget.clear()
        chaindata={}
        for item in self.emceeConfIntervalWidget.userDefinedParamTreeWidget.selectedItems():
            if ':chain:' in item.text(0):
                key,i=item.text(0).split(':chain:')
                try:
                    chaindata[key].append(int(i))
                except:
                    chaindata[key]=[int(i)]
        NRows = len(chaindata.keys())
        if NRows>0:
            ax={}
            firstkey=list(chaindata.keys())[0]
            for j,key in enumerate(chaindata.keys()):
                try:
                    ax[key]=self.emceeConfIntervalWidget.userDefinedChainMPLWidget.fig.add_subplot(NRows, 1, j+1, sharex=ax[firstkey])
                except:
                    ax[key] = self.emceeConfIntervalWidget.userDefinedChainMPLWidget.fig.add_subplot(NRows, 1, j+1)
                for i in chaindata[key]:
                    ax[key].plot(self.param_chain[key][i],'-')
                ax[key].set_xlabel('MC steps')
                ax[key].set_ylabel(key)
            self.emceeConfIntervalWidget.userDefinedChainMPLWidget.draw()
            self.emceeConfIntervalWidget.tabWidget.setCurrentIndex(1)

    def saveParameterError(self):
        fname=QFileDialog.getSaveFileName(caption='Save Parameter Errors as',filter='Parameter Error files (*.perr)',directory=self.curDir)[0]
        if os.path.splitext(fname)[1]=='':
            fname=fname+'.perr'
        text=self.emceeConfIntervalWidget.confIntervalTextEdit.toPlainText()
        fh=open(fname,'w')
        fh.write('# File save on %s\n'%time.asctime())
        fh.write('# Error calculated using MCMC Method\n')
        fh.write(text)
        fh.close()


    def undoFit(self):
        try:
            self.sfitParamTableWidget.cellChanged.disconnect()
            for i in range(self.mfitParamTabWidget.count()):
                mkey=self.mfitParamTabWidget.tabText(i)
                self.mfitParamTableWidget[mkey].cellChanged.disconnect()
        except:
            pass
        for row in range(self.sfitParamTableWidget.rowCount()):
            key=self.sfitParamTableWidget.item(row,0).text()
            self.sfitParamTableWidget.item(row,1).setText(self.format%(self.oldParams[key]))
            self.sfitParamTableWidget.item(row,1).setToolTip((key+' = '+self.format+' \u00B1 '+self.format)% (self.oldParams[key], 0.0))
        if self.fit.params['__mpar__']!={}:
            for i in range(self.mfitParamTabWidget.count()):
                mkey=self.mfitParamTabWidget.tabText(i)
                for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                    for col in range(1,self.mfitParamTableWidget[mkey].columnCount()):
                        parkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                        key='__%s_%s_%03d'%(mkey,parkey,row)
                        self.mfitParamTableWidget[mkey].item(row,col).setText(self.format%(self.oldmpar[mkey][parkey][row]))
                        self.mfitParamTableWidget[mkey].item(row, col).setToolTip((key+' = '+self.format+' \u00B1 '+self.format) % \
                                                                        (self.oldmpar[mkey][parkey][row], 0.0))
            #self.mfitParamData=copy.copy(self.oldmpar)
        self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
        for i in range(self.mfitParamTabWidget.count()):
            mkey = self.mfitParamTabWidget.tabText(i)
            self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
        self.update_plot()


        
        
    def addData(self,fnames=None):
        """
        fnames        :List of filenames
        """
        if self.dataListWidget.count()==0:
            self.fileNumber=0
        try:
            self.dataListWidget.itemSelectionChanged.disconnect()
        except:
            pass
        #try:
        if fnames is None:
            fnames,_=QFileDialog.getOpenFileNames(self,caption='Open data files',directory=self.curDir,\
                                                  filter='Data files (*.txt *.dat *.chi *.rrf)')
        if len(fnames)!=0:
            self.curDir=os.path.dirname(fnames[0])
            for fname in fnames:
                data_key=str(self.fileNumber)+'<>'+fname
                data_dlg=Data_Dialog(fname=fname,parent=self)
                data_dlg.setModal(True)
                data_dlg.closePushButton.setText('Cancel')
                if len(fnames)>1:
                    data_dlg.accept()
                else:
                    data_dlg.exec_()
                if data_dlg.acceptData:
                    self.dlg_data[data_key]=data_dlg.data
                    self.plotColIndex[data_key]=data_dlg.plotColIndex
                    self.plotColors[data_key]=data_dlg.plotColors
                    self.data[data_key]=data_dlg.externalData
                    self.expressions[data_key]=data_dlg.expressions
                    for key in self.data[data_key].keys():
                        self.plotWidget.add_data(self.data[data_key][key]['x'],self.data[data_key][key]['y'],\
                                                 yerr=self.data[data_key][key]['yerr'],name='%d:%s'%(self.fileNumber,key),color=self.data[data_key][key]['color'])
                    self.dataListWidget.addItem(data_key)
                    self.fileNames[self.fileNumber]=fname
                    self.fileNumber+=1
            #     else:
            #         QMessageBox.warning(self,'Import Error','Data file has been imported before.\
            #          Please remove the data file before importing again')
            # #except:
            # #    QMessageBox.warning(self,'File error','The file(s) do(es) not look like a data file. Please format it in x,y[,yerr] column format',QMessageBox.Ok)
        self.dataListWidget.clearSelection()
        self.dataListWidget.itemSelectionChanged.connect(self.dataFileSelectionChanged)
        self.dataListWidget.setCurrentRow(self.fileNumber-1)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

                
        
    def removeData(self):
        """
        """
        try:
            self.dataListWidget.itemSelectionChanged.disconnect()
        except:
            pass
        for item in self.dataListWidget.selectedItems():
            fnum,fname=item.text().split('<>')
            self.dataListWidget.takeItem(self.dataListWidget.row(item))
            for key in self.data[item.text()].keys():
                self.plotWidget.remove_data(['%s:%s'%(fnum,key)])
            del self.data[item.text()]
            del self.expressions[item.text()]
            del self.plotColIndex[item.text()]
            del self.plotColors[item.text()]
            del self.dlg_data[item.text()]

        if self.dataListWidget.count()>0:
            self.dataFileSelectionChanged()
        else:
            self.pfnames=[]
        self.dataListWidget.itemSelectionChanged.connect(self.dataFileSelectionChanged)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

            
        
        
    def create_paramDock(self):
        self.parSplitter=QSplitter(Qt.Vertical)
                
        self.fixedparamLayoutWidget=pg.LayoutWidget(self)
        
        xlabel=QLabel('x')
        self.fixedparamLayoutWidget.addWidget(xlabel)
        self.xLineEdit=QLineEdit('np.linspace(0.001,1,100)')
        self.fixedparamLayoutWidget.addWidget(self.xLineEdit,col=1)
        self.saveSimulatedButton=QPushButton("Save Simulated Curve")
        self.saveSimulatedButton.setEnabled(False)
        self.saveSimulatedButton.clicked.connect(self.saveSimulatedCurve)
        self.fixedparamLayoutWidget.addWidget(self.saveSimulatedButton,col=2)

        self.fixedparamLayoutWidget.nextRow()
        self.saveParamButton = QPushButton('Save Parameters')
        self.saveParamButton.clicked.connect(self.saveParameters)
        self.fixedparamLayoutWidget.addWidget(self.saveParamButton,col=1)
        self.loadParamButton = QPushButton('Load Parameters')
        self.loadParamButton.clicked.connect(lambda x: self.loadParameters(fname=None))
        self.fixedparamLayoutWidget.addWidget(self.loadParamButton, col=2)
        
        self.fixedparamLayoutWidget.nextRow()
        fixedParamLabel=QLabel('Fixed Parameters')
        self.fixedparamLayoutWidget.addWidget(fixedParamLabel, colspan=2)
        self.autoCalcCheckBox=QCheckBox('Auto Calculate')
        self.fixedparamLayoutWidget.addWidget(self.autoCalcCheckBox)
        self.autoCalcCheckBox.setChecked(True)
        self.autoCalcCheckBox.stateChanged.connect(self.changeAutoCalc)

        self.fixedparamLayoutWidget.nextRow()
        self.fixedParamTableWidget=pg.TableWidget()
        self.fixedParamTableWidget.setSizePolicy(QSizePolicy.Expanding,QSizePolicy.Expanding)
        self.fixedParamTableWidget.setEditable(editable=True)
        self.fixedParamTableWidget.setSizeAdjustPolicy(QAbstractScrollArea.AdjustToContents)
        self.fixedparamLayoutWidget.addWidget(self.fixedParamTableWidget,colspan=3)
        
        self.parSplitter.addWidget(self.fixedparamLayoutWidget)
        
        self.sfitparamLayoutWidget=pg.LayoutWidget()
        sfitParamLabel=QLabel('Single fitting parameters')
        self.sfitparamLayoutWidget.addWidget(sfitParamLabel)
        
        self.sfitparamLayoutWidget.nextRow()
        self.sfitParamTableWidget=pg.TableWidget()
        self.sfitParamTableWidget.setEditable(editable=True)
        self.sfitParamTableWidget.setSizePolicy(QSizePolicy.Expanding,QSizePolicy.Expanding)
        self.sfitParamTableWidget.setSizeAdjustPolicy(QAbstractScrollArea.AdjustToContents)
        #self.sfitParamTableWidget.cellDoubleClicked.connect(self.editFitParam)
        self.sfitparamLayoutWidget.addWidget(self.sfitParamTableWidget,colspan=3)
        self.sfitparamLayoutWidget.nextRow()
        self.sfitLabel=QLabel('')
        self.sfitSlider=QSlider(Qt.Horizontal)
        self.sfitSlider.setMinimum(1)
        self.sfitSlider.setMaximum(1000)
        self.sfitSlider.setSingleStep(10)
        self.sfitSlider.setTickInterval(10)
        self.sfitSlider.setValue(500)
        self.sfitparamLayoutWidget.addWidget(self.sfitLabel,col=0,colspan=1)
        self.sfitparamLayoutWidget.addWidget(self.sfitSlider,col=1,colspan=2)
        self.sfitParamTableWidget.cellClicked.connect(self.update_sfitSlider)

        self.parSplitter.addWidget(self.sfitparamLayoutWidget)
        
        self.mfitparamLayoutWidget=pg.LayoutWidget()
        mfitParamLabel=QLabel('Mutiple fitting parameters')
        self.mfitparamLayoutWidget.addWidget(mfitParamLabel,col=0, colspan=3)

        self.mfitparamLayoutWidget.nextRow()
        self.mfitParamCoupledCheckBox=QCheckBox('Coupled')
        self.mfitParamCoupledCheckBox.setEnabled(False)
        self.mfitParamCoupledCheckBox.stateChanged.connect(self.mfitParamCoupledCheckBoxChanged)
        self.mfitparamLayoutWidget.addWidget(self.mfitParamCoupledCheckBox,col=0)
        self.add_mpar_button=QPushButton('Add')
        self.add_mpar_button.clicked.connect(self.add_mpar)
        self.add_mpar_button.setDisabled(True)
        self.mfitparamLayoutWidget.addWidget(self.add_mpar_button,col=1)
        self.remove_mpar_button=QPushButton('Remove')
        self.mfitparamLayoutWidget.addWidget(self.remove_mpar_button,col=2)      
        self.remove_mpar_button.clicked.connect(self.remove_mpar)
        self.remove_mpar_button.setDisabled(True)
        
        self.mfitparamLayoutWidget.nextRow()
        self.mfitParamTabWidget=QTabWidget()
        self.mfitParamTabWidget.currentChanged.connect(self.mfitParamTabChanged)
        # self.mfitParamTableWidget=pg.TableWidget(sortable=False)
        # self.mfitParamTableWidget.cellDoubleClicked.connect(self.mparDoubleClicked)
        # self.mfitParamTableWidget.setEditable(editable=True)
        # self.mfitParamTableWidget.setSizePolicy(QSizePolicy.Expanding,QSizePolicy.Expanding)
        # self.mfitParamTableWidget.setSizeAdjustPolicy(QAbstractScrollArea.AdjustToContents)
        # #self.sfitParamTableWidget.cellDoubleClicked.connect(self.editFitParam)
        # self.mfitparamLayoutWidget.addWidget(self.mfitParamTableWidget,colspan=3)
        self.mfitparamLayoutWidget.addWidget(self.mfitParamTabWidget,colspan=3)
        self.mfitparamLayoutWidget.nextRow()
        self.mfitLabel=QLabel('')
        self.mfitSlider=QSlider(Qt.Horizontal)
        self.mfitSlider.setMinimum(1)
        self.mfitSlider.setSingleStep(10)
        self.mfitSlider.setTickInterval(10)
        self.mfitSlider.setMaximum(1000)
        self.mfitSlider.setValue(500)
        self.mfitparamLayoutWidget.addWidget(self.mfitLabel,col=0,colspan=1)
        self.mfitparamLayoutWidget.addWidget(self.mfitSlider,col=1,colspan=2)
        # self.mfitParamTableWidget.cellClicked.connect(self.update_mfitSlider)
        
        # self.mfitparamLayoutWidget.nextRow()
        # self.saveParamButton=QPushButton('Save Parameters')
        # self.saveParamButton.clicked.connect(self.saveParameters)
        # self.mfitparamLayoutWidget.addWidget(self.saveParamButton,col=1)
        # self.loadParamButton=QPushButton('Load Parameters')
        # self.loadParamButton.clicked.connect(lambda x: self.loadParameters(fname=None))
        # self.mfitparamLayoutWidget.addWidget(self.loadParamButton,col=2)
        self.parSplitter.addWidget(self.mfitparamLayoutWidget)

        self.genparamLayoutWidget=pg.LayoutWidget()
        genParameters=QLabel('Generated Parameters')
        self.autoSaveGenParamCheckBox=QCheckBox('Auto Save Selected')
        self.genparamLayoutWidget.addWidget(genParameters,col=0,colspan=2)
        self.genparamLayoutWidget.addWidget(self.autoSaveGenParamCheckBox,col=1)
        self.genparamLayoutWidget.nextRow()
        self.genParamListWidget=QListWidget()
        self.genParamListWidget.setSelectionMode(QAbstractItemView.ExtendedSelection)
        self.genParamListWidget.itemSelectionChanged.connect(self.plot_extra_param)
        self.genParamListWidget.itemDoubleClicked.connect(self.extra_param_doubleClicked)
        #self.genParamListWidget.setSizePolicy(QSizePolicy.Expanding,QSizePolicy.Expanding)
        self.genparamLayoutWidget.addWidget(self.genParamListWidget,colspan=2)
        self.genparamLayoutWidget.nextRow()
        self.saveGenParamButton=QPushButton('Save Generated Parameters')
        self.saveGenParamButton.clicked.connect(lambda x:self.saveGenParameters(bfname=None))
        self.genparamLayoutWidget.addWidget(self.saveGenParamButton,colspan=2)

        self.parSplitter.addWidget(self.genparamLayoutWidget)
        
        self.paramDock.addWidget(self.parSplitter)


    def changeAutoCalc(self):
        if self.autoCalcCheckBox.isChecked():
            self.autoCalculate=True
            self.update_plot()
        else:
            self.autoCalculate=False

    def mfitParamTabChanged(self,index):
        self.mkey=self.mfitParamTabWidget.tabText(index)
        if self.mkey!='':
            if self.mfitParamTableWidget[self.mkey].rowCount()==self.mpar_N[self.mkey]:
                self.remove_mpar_button.setDisabled(True)
            else:
                self.remove_mpar_button.setEnabled(True)


    def update_sfitSlider(self,row,col):
        if col==1:
            try:
                self.sfitSlider.valueChanged.disconnect()
                self.sfitSlider.sliderReleased.disconnect()
            except:
                pass
            key=self.sfitParamTableWidget.item(row,0).text()
            self.sfitLabel.setText(key)
            self.current_sfit_row=row
            value=self.fit.fit_params[key].value
            self.sfitSlider.setValue(500)
            self.sfitSlider.valueChanged.connect(self.sfitSliderChanged)
            self.sfitSlider.sliderReleased.connect(self.sfitSliderReleased)

    def sfitSliderChanged(self,value):
        if not self.sfitSlider.isSliderDown():
            self.sfitSlider.setDisabled(True)
            key=self.sfitParamTableWidget.item(self.current_sfit_row,0).text()
            pvalue=self.fit.fit_params[key].value+self.fit.fit_params[key].brute_step*(value-500)/500
            self.sfitParamTableWidget.item(self.current_sfit_row,1).setText(self.format%pvalue)
            QApplication.processEvents()
            self.sfitSlider.setEnabled(True)

    def sfitSliderReleased(self):
        key=self.sfitParamTableWidget.item(self.current_sfit_row,0).text()
        pvalue=self.fit.fit_params[key].value*(1+0.2*(self.sfitSlider.value()-500)/500)
        self.sfitParamTableWidget.item(self.current_sfit_row,1).setText(self.format%pvalue)
        QApplication.processEvents()

    def update_mfitSlider(self,row,col):
        if col!=0:
            try:
                self.mfitSlider.valueChanged.disconnect()
                self.mfitSlider.sliderReleased.disconnect()
            except:
                pass
            pkey = self.mfitParamTableWidget[self.mkey].horizontalHeaderItem(col).text()
            txt = self.mfitParamTableWidget[self.mkey].item(row, col).text()
            key = '__%s_%s_%03d' % (self.mkey, pkey, row)
            self.mfitLabel.setText(key)
            self.current_mfit_row=row
            self.current_mfit_col=col
            value=self.fit.fit_params[key].value
            self.mfitSlider.setValue(500)
            self.mfitSlider.valueChanged.connect(self.mfitSliderChanged)
            self.mfitSlider.sliderReleased.connect(self.mfitSliderReleased)

    def mfitSliderChanged(self,value):
        if not self.mfitSlider.isSliderDown():
            self.mfitSlider.setDisabled(True)
            pkey = self.mfitParamTableWidget[self.mkey].horizontalHeaderItem(self.current_mfit_col).text()
            txt = self.mfitParamTableWidget[self.mkey].item(self.current_mfit_row, self.current_mfit_col).text()
            key = '__%s_%s_%03d' % (self.mkey, pkey, self.current_mfit_row)
            pvalue=self.fit.fit_params[key].value+self.fit.fit_params[key].brute_step*(value-500)/500
            self.mfitParamTableWidget[self.mkey].item(self.current_mfit_row,self.current_mfit_col).setText(self.format%pvalue)
            QApplication.processEvents()
            self.mfitSlider.setEnabled(True)

    def mfitSliderReleased(self):
        pkey = self.mfitParamTableWidget[self.mkey].horizontalHeaderItem(self.current_mfit_col).text()
        txt = self.mfitParamTableWidget[self.mkey].item(self.current_mfit_row, self.current_mfit_col).text()
        key = '__%s_%s_%03d' % (self.mkey, pkey, self.current_mfit_row)
        pvalue = self.fit.fit_params[key].value * (1 + 0.2 * (self.mfitSlider.value() - 500) / 500)
        self.mfitParamTableWidget[self.mkey].item(self.current_mfit_row, self.current_mfit_col).setText(self.format % pvalue)
        QApplication.processEvents()


    def saveSimulatedCurve(self):
        """
        Saves the simulated curve in a user-supplied ascii file
        :return:
        """
        fname=QFileDialog.getSaveFileName(caption='Save As',filter='Text files (*.dat *.txt)',directory=self.curDir)[0]
        if fname!='':
            header='Simulated curve generated on %s\n'%time.asctime()
            header+='Category:%s\n'%self.curr_category
            header+='Function:%s\n'%self.funcListWidget.currentItem().text()
            for i in range(self.fixedParamTableWidget.rowCount()):
                header += '%s=%s\n' % (
                self.fixedParamTableWidget.item(i, 0).text(), self.fixedParamTableWidget.item(i, 1).text())
            for i in range(self.sfitParamTableWidget.rowCount()):
                header += '%s=%s\n' % (
                self.sfitParamTableWidget.item(i, 0).text(), self.sfitParamTableWidget.item(i, 1).text())
            for i in range(self.mfitParamTabWidget.count()):
                mkey = self.mfitParamTabWidget.tabText(i)
                for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                    vartxt = mkey+'_'+self.mfitParamTableWidget[mkey].item(row, 0).text()
                    for col in range(1, self.mfitParamTableWidget[mkey].columnCount()):
                        header += '%s_%s=%s\n' % (vartxt, self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text(),
                                              self.mfitParamTableWidget[mkey].item(row, col).text())
            if type(self.fit.x)==dict:
                text='col_names=[\'q\','
                keys=list(self.fit.x.keys())
                data=self.fit.x[keys[0]]
                for key in keys:
                    text+='\''+key+'\','
                    data=np.vstack((data,self.fit.yfit[key]))
                header+=text[:-1]+']\n'
                np.savetxt(fname,data.T,header=header,comments='#')
            else:
                header+='col_names=[\'q\',\'I\']'
                np.savetxt(fname,np.vstack((self.fit.x,self.fit.yfit)).T,header=header,comments='#')
        else:
            pass

        
    def mparDoubleClicked(self,row,col):
        mkey=self.mfitParamTabWidget.tabText(self.mfitParamTabWidget.currentIndex())
        if col!=0:
            try:
                self.mfitParamTableWidget[mkey].cellChanged.disconnect()
            except:
                pass
            pkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
            key='__%s_%s_%03d'%(mkey,pkey,row)
            ovalue=self.fit.fit_params[key].value
            ovary=self.fit.fit_params[key].vary
            ominimum=self.fit.fit_params[key].min
            omaximum=self.fit.fit_params[key].max
            oexpr=self.fit.fit_params[key].expr
            obrute_step=self.fit.fit_params[key].brute_step
            dlg=minMaxDialog(ovalue,vary=ovary,minimum=ominimum,maximum=omaximum,expr=oexpr,brute_step=obrute_step,title=key)
            if dlg.exec_():
                value,vary,maximum,minimum,expr,brute_step=(dlg.value,dlg.vary,dlg.maximum,dlg.minimum,dlg.expr,dlg.brute_step)
            else:
                value,vary,maximum,minimum,expr,brute_step=copy.copy(ovalue),copy.copy(ovary),copy.copy(omaximum),copy.copy(ominimum),copy.copy(oexpr),copy.copy(obrute_step)
            self.mfitParamTableWidget[mkey].item(row,col).setText(self.format%value)
            if vary:
                self.mfitParamTableWidget[mkey].item(row, col).setCheckState(Qt.Checked)
            else:
                self.mfitParamTableWidget[mkey].item(row, col).setCheckState(Qt.Unchecked)

            try:
                self.mfitParamData[mkey][pkey][row] = value
                # self.fit.fit_params[key].set(value=value)
                if expr == 'None':
                    expr = ''
                self.fit.fit_params[key].set(value=value, vary=vary, min=minimum, max=maximum, expr=expr,
                                             brute_step=brute_step)
                self.update_plot()
            except:
                self.mfitParamTableWidget[mkey].item(row, col).setText(self.format % ovalue)
                self.mfitParamData[mkey][pkey][row] = ovalue
                self.fit.fit_params[key].set(value=ovalue, vary=ovary, min=ominimum, max=omaximum, expr=oexpr,
                                             brute_step=brute_step)
                self.update_plot()
                QMessageBox.warning(self,'Parameter Error','Some parameter value you just entered are not correct. Please enter the values carefully',QMessageBox.Ok)

            self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
            self.errorAvailable = False
            self.reuse_sampler = False
            self.calcConfInterButton.setDisabled(True)


    def mfitParamCoupledCheckBoxChanged(self):
        if self.mfitParamCoupledCheckBox.isChecked() and self.mfitParamTabWidget.count()>1:
            mparRowCounts=[self.mfitParamTableWidget[self.mfitParamTabWidget.tabText(i)].rowCount() for i in range(self.mfitParamTabWidget.count())]
            if not all(x == mparRowCounts[0] for x in mparRowCounts):
                cur_index=self.mfitParamTabWidget.currentIndex()
                cur_key=self.mfitParamTabWidget.tabText(cur_index)
                for i in range(self.mfitParamTabWidget.count()):
                    if i != cur_index:
                        mkey=self.mfitParamTabWidget.tabText(i)
                        try:
                            self.mfitParamTableWidget[mkey].cellChanged.disconnect()
                        except:
                            pass
                        rowCount=self.mfitParamTableWidget[mkey].rowCount()
                        self.mfitParamTabWidget.setCurrentIndex(i)
                        if rowCount>mparRowCounts[cur_index]:
                            self.mfitParamTableWidget[mkey].clearSelection()
                            self.mfitParamTableWidget[mkey].setRangeSelected(
                                QTableWidgetSelectionRange(mparRowCounts[cur_index],0,rowCount-1,0),True)
                            self.remove_uncoupled_mpar()
                        elif rowCount<mparRowCounts[cur_index]:
                            for j in range(rowCount,mparRowCounts[cur_index]):
                                self.mfitParamTableWidget[mkey].clearSelection()
                                self.mfitParamTableWidget[mkey].setCurrentCell(j-1,0)
                                self.add_uncoupled_mpar()
                        self.mfitParamTableWidget[mkey].setSelectionBehavior(QAbstractItemView.SelectItems)
                self.mfitParamTabWidget.setCurrentIndex(cur_index)
                self.errorAvailable = False
                self.reuse_sampler = False
                self.calcConfInterButton.setDisabled(True)

    def add_mpar(self):
        if self.mfitParamCoupledCheckBox.isChecked() and self.mfitParamTabWidget.count()>1:
            self.add_coupled_mpar()
        else:
            self.add_uncoupled_mpar()
        self.update_plot()
        self.remove_mpar_button.setEnabled(True)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

    def remove_mpar(self):
        if self.mfitParamCoupledCheckBox.isChecked() and self.mfitParamTabWidget.count()>1:
            self.remove_coupled_mpar()
        else:
            self.remove_uncoupled_mpar()
        self.update_plot()
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

    def add_coupled_mpar(self):
        cur_index=self.mfitParamTabWidget.currentIndex()
        mkey = self.mfitParamTabWidget.tabText(cur_index)
        if len(self.mfitParamTableWidget[mkey].selectedItems())!=0:
            curRow=self.mfitParamTableWidget[mkey].currentRow()
            for i in range(self.mfitParamTabWidget.count()):
                self.mfitParamTabWidget.setCurrentIndex(i)
                tkey=self.mfitParamTabWidget.tabText(i)
                self.mfitParamTableWidget[tkey].clearSelection()
                self.mfitParamTableWidget[tkey].setCurrentCell(curRow,0)
                self.add_uncoupled_mpar()
        self.mfitParamTabWidget.setCurrentIndex(cur_index)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)

    def remove_coupled_mpar(self):
        cur_index=self.mfitParamTabWidget.currentIndex()
        mkey = self.mfitParamTabWidget.tabText(cur_index)
        selRows = list(set([item.row() for item in self.mfitParamTableWidget[mkey].selectedItems()]))
        if len(selRows) != 0:
            for i in range(self.mfitParamTabWidget.count()):
                self.mfitParamTabWidget.setCurrentIndex(i)
                tkey=self.mfitParamTabWidget.tabText(i)
                self.mfitParamTableWidget[tkey].clearSelection()
                self.mfitParamTableWidget[tkey].setRangeSelected(
                    QTableWidgetSelectionRange(selRows[0], 0, selRows[-1], 0), True)
                self.remove_uncoupled_mpar()
        self.mfitParamTabWidget.setCurrentIndex(cur_index)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
        
    def add_uncoupled_mpar(self):

        cur_index = self.mfitParamTabWidget.currentIndex()
        mkey=self.mfitParamTabWidget.tabText(self.mfitParamTabWidget.currentIndex())
        try:
            self.mfitParamTableWidget[mkey].cellChanged.disconnect()
        except:
            pass
        NCols=self.mfitParamTableWidget[mkey].columnCount()
        if len(self.mfitParamTableWidget[mkey].selectedItems())!=0:
            curRow=self.mfitParamTableWidget[mkey].currentRow()
            #if curRow!=0:
            self.mfitParamTableWidget[mkey].insertRow(curRow)
            self.mfitParamTableWidget[mkey].setRow(curRow,self.mfitParamData[mkey][curRow])
            self.mfitParamData[mkey]=np.insert(self.mfitParamData[mkey],curRow,self.mfitParamData[mkey][curRow],0)
            NRows = self.mfitParamTableWidget[mkey].rowCount()
            for col in range(NCols):
                pkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                if col!=0:
                    for row in range(NRows-1, curRow,-1):
                        key='__%s_%s_%03d'%(mkey, pkey,row)
                        nkey = '__%s_%s_%03d' % (mkey,pkey,row-1)
                        if key in self.fit.fit_params.keys():
                            val,vary,min,max,expr,bs = self.mfitParamData[mkey][row][col],self.fit.fit_params[nkey].vary, \
                                                      self.fit.fit_params[nkey].min,self.fit.fit_params[nkey].max, \
                                                      self.fit.fit_params[nkey].expr,self.fit.fit_params[nkey].brute_step
                            self.fit.fit_params[key].set(value=val,vary=vary,min=min,max=max,expr=expr,brute_step=bs)
                        else:
                            val,vary,min,max,expr,bs=self.mfitParamData[mkey][row][col],self.fit.fit_params[nkey].vary,self.fit.fit_params[nkey].min, \
                                                 self.fit.fit_params[nkey].max,self.fit.fit_params[nkey].expr, \
                                                 self.fit.fit_params[nkey].brute_step
                            self.fit.fit_params.add(key,value=val,vary=vary,min=min,max=max,expr=expr,brute_step=bs)
                        item=self.mfitParamTableWidget[mkey].item(row,col)
                        item.setText(self.format%val)
                        item.setFlags(
                            Qt.ItemIsUserCheckable | Qt.ItemIsEnabled | Qt.ItemIsEditable | Qt.ItemIsSelectable)
                        if self.fit.fit_params[key].vary > 0:
                            item.setCheckState(Qt.Checked)
                        else:
                            item.setCheckState(Qt.Unchecked)
                        item.setToolTip((key+' = '+self.format+' \u00B1 '+self.format) % \
                                                                (self.fit.fit_params[key].value, 0.0))
                    # This is to make the newly inserted row checkable
                    item = self.mfitParamTableWidget[mkey].item(curRow, col)
                    item.setFlags(Qt.ItemIsUserCheckable | Qt.ItemIsEnabled | Qt.ItemIsEditable | Qt.ItemIsSelectable)
                    key = '__%s_%s_%03d' % (mkey, pkey, curRow)
                    item.setText(self.format%self.fit.fit_params[key].value)
                    item.setToolTip((key + ' = ' + self.format + ' \u00B1 ' + self.format) % \
                                    (self.fit.fit_params[key].value, 0.0))
                    if self.fit.fit_params[key].vary>0:
                        item.setCheckState(Qt.Checked)
                    else:
                        item.setCheckState(Qt.Unchecked)
                else:
                    item = self.mfitParamTableWidget[mkey].item(curRow, col)
                    item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsEditable | Qt.ItemIsSelectable)
                self.fit.params['__mpar__'][mkey][pkey].insert(curRow, self.mfitParamData[mkey][curRow][col])
            self.update_mfit_parameters_new()
            self.update_plot()
            self.errorAvailable = False
            self.reuse_sampler = False
            self.calcConfInterButton.setDisabled(True)
            # self.remove_mpar_button.setEnabled(True)
            self.mfitParamTabWidget.setCurrentIndex(cur_index)
        else:
            QMessageBox.warning(self,'Warning','Please select a row at which you would like to add a set of parameters',QMessageBox.Ok)
        self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
            
    def remove_uncoupled_mpar(self):
        mkey = self.mfitParamTabWidget.tabText(self.mfitParamTabWidget.currentIndex())
        selrows=list(set([item.row() for item in self.mfitParamTableWidget[mkey].selectedItems()]))
        num=self.mfitParamTableWidget[mkey].rowCount()-len(selrows)
        if num<self.mpar_N[mkey]:
            QMessageBox.warning(self,'Selection error','The minimum number of rows required for this function to work is %d.\
             You can only remove %d rows'%(self.mpar_N[mkey],num),QMessageBox.Ok)
            return
        # if self.mfitParamTableWidget[mkey].rowCount()-1 in selrows:
        #     QMessageBox.warning(self, 'Selection error',
        #                         'Cannot remove the last row. Please select the rows other than the last row', QMessageBox.Ok)
        #     return
        try:
            self.mfitParamTableWidget[mkey].cellChanged.disconnect()
        except:
            pass
        if selrows!=[]:
            selrows.sort(reverse=True)
            for row in selrows:
                maxrow=self.mfitParamTableWidget[mkey].rowCount()
                for trow in range(row,maxrow):
                    for col in range(self.mfitParamTableWidget[mkey].columnCount()):
                        pkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                        if trow<maxrow-1:
                            key1='__%s_%s_%03d'%(mkey,pkey,trow)
                            key2='__%s_%s_%03d'%(mkey,pkey,trow+1)
                            self.fit.params['__mpar__'][mkey][pkey][trow] = copy.copy(self.fit.params['__mpar__'][mkey][pkey][trow + 1])
                            if col!=0:
                                self.fit.fit_params[key1]=copy.copy(self.fit.fit_params[key2])
                                del self.fit.fit_params[key2]
                        else:
                            key1='__%s_%s_%03d'%(mkey,pkey,trow)
                            # if col!=0:
                            del self.fit.params['__mpar__'][mkey][pkey][trow]
                                # del self.fit.fit_params[key1]
                self.mfitParamTableWidget[mkey].removeRow(row)
                self.mfitParamData[mkey]=np.delete(self.mfitParamData[mkey],row,axis=0)
            #updating the tooltips after removal of rows
            for col in range(1,self.mfitParamTableWidget[mkey].columnCount()):
                pkey = self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                    item=self.mfitParamTableWidget[mkey].item(row, col)
                    key = '__%s_%s_%03d' % (mkey, pkey, row)
                    item.setToolTip((key + ' = ' + self.format + ' \u00B1 ' + self.format) % \
                        (self.fit.fit_params[key].value, 0.0))
        else:
            QMessageBox.warning(self,'Nothing selected','No item is selected for removal',QMessageBox.Ok)
        self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
        self.fit.func.output_params={'scaler_parameters': {}}
        self.update_plot()
        if self.mfitParamTableWidget[mkey].rowCount()==self.mpar_N[mkey]:
            self.remove_mpar_button.setDisabled(True)
        self.errorAvailable = False
        self.reuse_sampler = False
        self.calcConfInterButton.setDisabled(True)
            
        
    def saveGenParameters(self,bfname=None):
        # if len(self.genParamListWidget.selectedItems())==1:
        if bfname is None:
            bfname = QFileDialog.getSaveFileName(self, 'Provide the prefix of the generated files',self.curDir)[0]
        if bfname!='':
            bfname=os.path.splitext(bfname)[0]
        else:
            return
        selParams=self.genParamListWidget.selectedItems()
        for params in selParams:
            text=params.text()
            parname,var=text.split(' : ')
            fname=bfname+'_'+parname+'.txt'
            # if fname!='':
            #     if fname[-4:]!='.txt':
            #         fname=fname+'.txt'
            header='Generated output file on %s\n'%time.asctime()
            header += 'Category=%s\n' % self.curr_category
            header += 'Function=%s\n' % self.funcListWidget.currentItem().text()
            added_par=[]
            for i in range(self.fixedParamTableWidget.rowCount()):
                par, val = self.fixedParamTableWidget.item(i, 0).text(), self.fixedParamTableWidget.item(i, 1).text()
                if 'meta' in self.fit.params['output_params'][parname]:
                    if par in self.fit.params['output_params'][parname]['meta'].keys():
                        header += '%s=%s\n' % (par, str(self.fit.params['output_params'][parname]['meta'][par]))
                        added_par.append(par)
                    else:
                        header += '%s=%s\n' % (par, val)
                else:
                    header+='%s=%s\n'%(par,val)
            if 'meta' in self.fit.params['output_params'][parname]:
                for metakey in self.fit.params['output_params'][parname]['meta'].keys():
                    if metakey not in added_par:
                        header+='%s=%s\n'%(metakey,str(self.fit.params['output_params'][parname]['meta'][metakey]))
            # for fi in range(self.fixedParamTableWidget.rowCount()):
            #     par,val=self.fixedParamTableWidget.item(fi,0).text(),self.fixedParamTableWidget.item(fi,1).text()
            #     header+='%s=%s\n'%(par,val)
            for i in range(self.sfitParamTableWidget.rowCount()):
                par,val=self.sfitParamTableWidget.item(i,0).text(),self.sfitParamTableWidget.item(i,1).text()
                header+='%s=%s\n'%(par,val)
            for k in range(self.mfitParamTabWidget.count()):
                mkey=self.mfitParamTabWidget.tabText(k)
                for i in range(self.mfitParamTableWidget[mkey].rowCount()):
                    vartxt=self.mfitParamTableWidget[mkey].item(i,0).text()
                    for j in range(1,self.mfitParamTableWidget[mkey].columnCount()):
                        header+='%s_%s=%s\n'%(vartxt,self.mfitParamTableWidget[mkey].horizontalHeaderItem(j).text(),
                                              self.mfitParamTableWidget[mkey].item(i,j).text())

            if 'names' in self.fit.params['output_params'][parname]:
                header += "col_names=%s\n" % str(self.fit.params['output_params'][parname]['names'])
            else:
                lvar=eval(var)
                if 'meta' in lvar:
                    lvar.remove('meta')
                header += "col_names=%s\n" % str(lvar)

            header=header.encode("ascii","ignore")
            header=header.decode()
            if var=="['x', 'y', 'meta']" or var == "['x', 'y']":
                header+='x\ty\n'
                res=np.vstack((self.fit.params['output_params'][parname]['x'], self.fit.params['output_params'][parname]['y'])).T
                np.savetxt(fname,res,header=header,comments='#')
            elif var=="['x', 'y', 'yerr', 'meta']" or var=="['x', 'y', 'yerr']":
                header+='x\ty\tyerr\n'
                res=np.vstack((self.fit.params['output_params'][parname]['x'], self.fit.params['output_params'][parname]['y'],self.fit.params['output_params'][parname]['yerr'])).T
                np.savetxt(fname,res,header=header,comments='#')
            elif var=="['x', 'y', 'z', 'meta']" or var=="['x', 'y', 'z']":
                res=[]
                header+='x\ty\tz\n'
                for i in range(self.fit.params['output_params'][parname]['x'].shape[1]):
                    for j in range(self.fit.params['output_params'][parname]['x'].shape[0]):
                        res.append([self.fit.params['output_params'][parname][t][i,j] for t in ['x','y','z']])
                res=np.array(res)
                np.savetxt(fname,res,header=header,comments='#')
            else:
                QMessageBox.warning(self,'Format error','The data is in some different format and couldnot be saved.',QMessageBox.Ok)
        # else:
        #     QMessageBox.warning(self,'Selection Error','Please select a single generated data to be saved.',QMessageBox.Ok)
        
        
    def saveParameters(self):
        """
        Saves all the fixed and fitted parameteres in a file
        """
        fname=QFileDialog.getSaveFileName(self,caption='Save parameters as',directory=self.curDir,filter='Parameter files (*.par)')[0]
        if fname!='':
            if fname[-4:]!='.par':
                fname=fname+'.par'
            fh=open(fname,'w')
            fh.write('#File saved on %s\n'%time.asctime())
            fh.write('#Category: %s\n'%self.categoryListWidget.currentItem().text())
            fh.write('#Function: %s\n'%self.funcListWidget.currentItem().text())
            fh.write('#Xrange=%s\n'%self.xLineEdit.text())
            fh.write('#Fit Range=%s\n'%self.xminmaxLineEdit.text())
            fh.write('#Fit Method=%s\n'%self.fitMethodComboBox.currentText())
            fh.write('#Fit Scale=%s\n'%self.fitScaleComboBox.currentText())
            fh.write('#Fit Iterations=%s\n'%self.fitIterationLineEdit.text())
            fh.write('#Fixed Parameters:\n')
            fh.write('#param\tvalue\n')
            for row in range(self.fixedParamTableWidget.rowCount()):
                txt=self.fixedParamTableWidget.item(row,0).text()
                if txt in self.fit.params['choices'].keys():
                    fh.write(txt+'\t'+self.fixedParamTableWidget.cellWidget(row, 1).currentText()+'\n')
                else:
                    fh.write(txt+'\t'+self.fixedParamTableWidget.item(row,1).text()+'\n')
            fh.write('#Single fitting parameters:\n')
            fh.write('#param\tvalue\tfit\tmin\tmax\texpr\tbrute_step\n')
            for row in range(self.sfitParamTableWidget.rowCount()):
                parname=self.sfitParamTableWidget.item(row,0).text()
                par=self.sfitParamTableWidget.item(row,1)
                parval=par.text()
                if par.checkState()==Qt.Checked:
                    parfit='1'
                else:
                    parfit='0'
                parmin=self.sfitParamTableWidget.item(row,2).text()
                parmax=self.sfitParamTableWidget.item(row,3).text()
                parexpr=self.sfitParamTableWidget.item(row,4).text()
                parbrute=self.sfitParamTableWidget.item(row,5).text()
                fh.write('%s\t%s\t%s\t%s\t%s\t%s\t%s\n'%(parname,parval,parfit,parmin,parmax,parexpr,parbrute))
            if self.fit.params['__mpar__']!={}:
                fh.write('#Multiple fitting parameters:\n')
                fh.write('#param\tvalue\tfit\tmin\tmax\texpr\tbrute_step\n')
                for i in range(self.mfitParamTabWidget.count()):
                    mkey=self.mfitParamTabWidget.tabText(i)
                    for col in range(self.mfitParamTableWidget[mkey].columnCount()):
                        pkey = self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                        if col!=0:
                            for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                                parname='__%s_%s_%03d'%(mkey,pkey,row)
                                par=self.mfitParamTableWidget[mkey].item(row,col)
                                parval=par.text()
                                if par.checkState()==Qt.Checked:
                                    parfit='1'
                                else:
                                    parfit='0'
                                parmin=str(self.fit.fit_params[parname].min)
                                parmax=str(self.fit.fit_params[parname].max)
                                parexpr=str(self.fit.fit_params[parname].expr)
                                parbrute=str(self.fit.fit_params[parname].brute_step)
                                fh.write('%s\t%s\t%s\t%s\t%s\t%s\t%s\n'%(parname,parval,parfit,parmin,parmax,parexpr,parbrute))
                        else:
                            for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                                parname = '__%s_%s_%03d' % (mkey, pkey, row)
                                par = self.mfitParamTableWidget[mkey].item(row, col)
                                parval = par.text()
                                fh.write('%s\t%s\n' % (parname, parval))
            fh.close()
        
        
    def loadParameters(self,fname=None):
        """
        loads parameters from a parameter file
        """
        # if self.funcListWidget.currentItem() is not None:
        if fname is None:
            fname=QFileDialog.getOpenFileName(self,caption='Open parameter file',directory=self.curDir,filter='Parameter files (*.par)')[0]
        else:
            fname=fname
        if fname!='':
            self.curDir = os.path.dirname(fname)
            try:
                self.funcListWidget.itemSelectionChanged.disconnect()
            except:
                pass
            try:
                fh=open(fname,'r')
                lines=fh.readlines()
                category=lines[1].split(': ')[1].strip()
                cat_item=self.categoryListWidget.findItems(category,Qt.MatchExactly)
                self.categoryListWidget.setCurrentItem(cat_item[0])
                self.funcListWidget.clearSelection()
                func=lines[2].split(': ')[1].strip()
                func_item=self.funcListWidget.findItems(func,Qt.MatchExactly)
                self.funcListWidget.itemSelectionChanged.connect(self.functionChanged)
                self.funcListWidget.setCurrentItem(func_item[0])
                #self.fit.func.init_params()
                if func==self.funcListWidget.currentItem().text():
                    lnum=3
                    sfline=None
                    mfline=None
                    for line in lines[3:]:
                        if '#Xrange=' in line:
                            self.xLineEdit.setText(line.strip().split('=')[1])
                        elif '#Fit Range=' in line:
                            self.xminmaxLineEdit.setText(line.strip().split('=')[1])
                            fline=lnum+1
                        elif '#Fit Method=' in line:
                            self.fitMethodComboBox.setCurrentText(line.strip().split('=')[1])
                            fline=lnum+1
                        elif '#Fit Scale=' in line:
                            self.fitScaleComboBox.setCurrentText(line.strip().split('=')[1])
                            fline=lnum+1
                        elif '#Fit Iterations=' in line:
                            self.fitIterationLineEdit.setText(line.strip().split('=')[1])
                            fline=lnum+1
                        elif line=='#Fixed Parameters:\n':
                            fline=lnum+2
                        elif line=='#Single fitting parameters:\n':
                            sfline=lnum+2
                        elif line=='#Multiple fitting parameters:\n':
                            mfline=lnum+2
                        lnum+=1
                    if sfline is None:
                        sendnum=lnum
                    else:
                        sendnum=sfline-2
                    if mfline is None:
                        mendnum=lnum
                    else:
                        mendnum=mfline-2
                    for line in lines[fline:sendnum]:
                        key,val=line.strip().split('\t')
                        try:
                            val=eval(val.strip())
                        except:
                            val=val.strip()
                        self.fit.params[key]=val
                    if sfline is not None:
                        for line in lines[sfline:mendnum]:
                            parname,parval,parfit,parmin,parmax,parexpr,parbrute=line.strip().split('\t')
                            self.fit.params[parname]=float(parval)
                            self.fit.fit_params[parname].set(value=float(parval),vary=int(parfit),min=float(parmin),max=float(parmax))
                            try:
                                self.fit.fit_params[parname].set(expr=eval(parexpr))
                            except:
                                self.fit.fit_params[parname].set(expr=str(parexpr))
                            try:
                                self.fit.fit_params[parname].set(brute_step=eval(parbrute))
                            except:
                                self.fit.fit_params[parname].set(brute_step=str(parbrute))

                    if mfline is not None:
                        self.mfitParamCoupledCheckBox.setEnabled(True)
                        val={}
                        expr={}
                        pmin={}
                        pmax={}
                        pbrute={}
                        pfit={}
                        for line in lines[mfline:]:
                            tlist=line.strip().split('\t')
                            if len(tlist)>2:
                                parname,parval,parfit,parmin,parmax,parexpr,parbrute=tlist
                                val[parname]=float(parval)
                                pmin[parname]=float(parmin)
                                pmax[parname]=float(parmax)
                                pfit[parname]=int(parfit)
                                try:
                                    expr[parname]=eval(parexpr)
                                except:
                                    expr[parname]=str(parexpr)
                                try:
                                    pbrute[parname]=eval(parbrute)
                                except:
                                    pbrute[parname]=str(parbrute)
                                try: # Here the expr is set to None and will be taken care at the for loop just after this for loop
                                    self.fit.fit_params[parname].set(val[parname], vary=pfit[parname],
                                                                     min=pmin[parname],
                                                                     max=pmax[parname], expr=None,
                                                                     brute_step=pbrute[parname])
                                except:
                                    self.fit.fit_params.add(parname, value=val[parname], vary=pfit[parname],
                                                            min=pmin[parname],
                                                            max=pmax[parname], expr=None,
                                                            brute_step=pbrute[parname])
                                mkey, pkey, num = parname[2:].split('_')
                                num = int(num)
                                try:
                                    self.fit.params['__mpar__'][mkey][pkey][num] = float(parval)
                                except:
                                    self.fit.params['__mpar__'][mkey][pkey].insert(num, float(parval))
                            else:
                                parname,parval=tlist

                                mkey,pkey,num=parname[2:].split('_')
                                num=int(num)
                                try:
                                    self.fit.params['__mpar__'][mkey][pkey][num]=parval
                                except:
                                    self.fit.params['__mpar__'][mkey][pkey].insert(num,parval)

                        for parname in val.keys(): #Here is the expr is put into the parameters
                            try:
                                self.fit.fit_params[parname].set(value=val[parname], vary=pfit[parname], min=pmin[parname],
                                                                 max=pmax[parname], expr=expr[parname], brute_step=pbrute[parname])
                            except:
                                self.fit.fit_params.add(parname, value=val[parname], vary=pfit[parname], min=pmin[parname],
                                                        max=pmax[parname], expr=expr[parname], brute_step=pbrute[parname])
                    try:
                        self.fixedParamTableWidget.cellChanged.disconnect()
                        self.sfitParamTableWidget.cellChanged.disconnect()
                        for i in range(self.mfitParamTabWidget.count()):
                            mkey = self.mfitParamTabWidget.tabText(i)
                            self.mfitParamTableWidget[mkey].cellChanged.disconnect()
                    except:
                        pass
                    self.update_fixed_parameters()
                    self.update_fit_parameters()
                    self.fixedParamTableWidget.cellChanged.connect(self.fixedParamChanged)
                    self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
                    for i in range(self.mfitParamTabWidget.count()):
                        mkey=self.mfitParamTabWidget.tabText(i)
                        self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
                    if len(self.dataListWidget.selectedItems())>0:
                        self.xminmaxChanged()
                    else:
                        self.xChanged()
                    self.errorAvailable=False
                    self.reuse_sampler=False
                    self.calcConfInterButton.setDisabled(True)
                else:
                    QMessageBox.warning(self, 'File error',
                                        'This parameter file does not belong to function: %s' % func, QMessageBox.Ok)
            except:
                QMessageBox.warning(self,'File Import Error','Some problems in the parameter file\n'+traceback.format_exc(), QMessageBox.Ok)
        # else:
        #     QMessageBox.warning(self,'Function error','Please select a function first before loading parameter file.', QMessageBox.Ok)


        
    def create_plotDock(self):
        self.plotSplitter=QSplitter(Qt.Vertical)
        #self.plotLayoutWidget=pg.LayoutWidget(self)
        self.plotWidget=PlotWidget()
        self.plotWidget.setXLabel('x',fontsize=5)
        self.plotWidget.setYLabel('y',fontsize=5)
        self.plotSplitter.addWidget(self.plotWidget)

        self.extra_param_1DplotWidget=PlotWidget()
        self.extra_param_1DplotWidget.setXLabel('x',fontsize=5)
        self.extra_param_1DplotWidget.setYLabel('y',fontsize=5)
        self.plotSplitter.addWidget(self.extra_param_1DplotWidget)

        self.plotDock.addWidget(self.plotSplitter)

    def create_fitResultDock(self):
        self.fitResultsLayoutWidget = pg.LayoutWidget()
        fitResults = QLabel('Fit Results')
        self.fitResultsLayoutWidget.addWidget(fitResults, colspan=1)
        self.fitResultsLayoutWidget.nextRow()
        self.fitResultTextEdit = QTextEdit()
        self.fitResultsLayoutWidget.addWidget(self.fitResultTextEdit, colspan=1)
        self.fitResultDock.addWidget(self.fitResultsLayoutWidget)

    def update_catagories(self):
        """
        Reads all the modules in the the Functions directory and populates the funcListWidget
        """
        self.categoryListWidget.clear()
        self.categories=sorted([path for path in os.listdir('./Functions/') if path[:2]!='__' and os.path.isdir('./Functions/'+path)])
        #self.catagories=sorted([m.split('.')[0] for m in modules if m[:2]!='__'],key=str.lower)
        self.categoryListWidget.addItems(self.categories)

        
    def update_functions(self):
        """
        Depending upon the selected category this populates the funcListWidget
        """
        self.saveSimulatedButton.setEnabled(False)
        try:
            self.funcListWidget.itemSelectionChanged.disconnect()
            self.funcListWidget.itemDoubleClicked.disconnect()
        except:
            pass
        self.funcListWidget.clear()
        self.curr_category=self.categoryListWidget.currentItem().text()
        self.modules=[]
        for module in os.listdir('./Functions/'+self.curr_category):
            if module!='__init__.py' and module[-2:]=='py':
                self.modules.append(module[:-3])
        self.modules=sorted(self.modules,key=str.lower)
        self.funcListWidget.addItems(self.modules)
        for i in range(self.funcListWidget.count()):
            mname=self.funcListWidget.item(i).text()
            module='Functions.%s.%s'%(self.curr_category,mname)
            if module not in sys.modules:
                self.curr_funcClass[module]=import_module(module)
            else:
                self.curr_funcClass[module]=reload(self.curr_funcClass[module])
            self.funcListWidget.item(i).setToolTip(getattr(self.curr_funcClass[module],self.funcListWidget.item(i).text()).__init__.__doc__)
        self.funcListWidget.itemSelectionChanged.connect(self.functionChanged)
        self.funcListWidget.itemDoubleClicked.connect(self.openFunction)
        
    def functionChanged(self):
        if len(self.funcListWidget.selectedItems())<=1:
            self.sfitLabel.clear()
            self.mfitLabel.clear()
            self.sfitSlider.setValue(500)
            self.mfitSlider.setValue(500)
            self.gen_param_items=[]
            self.curr_module=self.funcListWidget.currentItem().text()
            module='Functions.%s.%s'%(self.curr_category,self.curr_module)
            self.mfitParamCoupledCheckBox.setEnabled(False)
            try:
                if module not in sys.modules:
                    self.curr_funcClass[module]=import_module(module)
                else:
                    self.curr_funcClass[module]=reload(self.curr_funcClass[module])
                mpath=os.path.join('Functions',self.curr_category,self.curr_module+'.py')
                fh=open(mpath,'r')
                lines=fh.readlines()
                for i,line in enumerate(lines):
                    if '__name__' in line:
                        lnum=i+1
                        break
                if 'x' in lines[lnum]:
                    self.xline=lines[lnum].split('=')[1].strip()
                else:
                    self.xline='np.linspace(0.0,1.0,100)'
                self.xLineEdit.setText(self.xline)
                self.fixedParamTableWidget.clear()
                self.sfitParamTableWidget.clear()
                self.mfitParamTabWidget.clear()
                # self.mfitParamTableWidget.clear()
                self.genParamListWidget.clear()
                self.fchanged = True
                self.update_parameters()
                self.saveSimulatedButton.setEnabled(True)
                self.errorAvailable = False
                self.reuse_sampler = False
                self.calcConfInterButton.setDisabled(True)
            except:
                QMessageBox.warning(self,'Function Error','Some syntax error in the function still exists.\n'+traceback.format_exc(),QMessageBox.Ok)
        else:
            QMessageBox.warning(self,'Function Error', 'Please select one function at a time', QMessageBox.Ok)
        
    def update_parameters(self):
        """
        Depending upon the selection of the function this updates the reloads the parameters required for the function
        """
        try:
            self.fixedParamTableWidget.cellChanged.disconnect()
            self.sfitParamTableWidget.cellChanged.disconnect()
            for i in range(self.mfitParamTabWidget.count()):
                mkey=self.mfitParamTabWidget.tabText(i)
                self.mfitParamTableWidget[mkey].cellChanged.disconnect()
        except:
            pass
        try:
            self.x=eval(self.xLineEdit.text())
        except:
            QMessageBox.warning(self,'Parameter Error','The value you just entered is not correct.\n'+traceback.format_exc(),QMessageBox.Ok)
        self.curr_module=self.funcListWidget.currentItem().text()
        module='Functions.%s.%s'%(self.curr_category,self.curr_module)
        self.fit=Fit(getattr(self.curr_funcClass[module],self.funcListWidget.currentItem().text()),self.x)
        if '__mpar__' in self.fit.params.keys() and self.fit.params['__mpar__'] != {}:
            self.mpar_keys = list(self.fit.params['__mpar__'].keys())
            pkey=list(self.fit.params['__mpar__'][self.mpar_keys[0]].keys())[0]
            self.mpar_N={}
            for mkey in self.mpar_keys:
                self.mpar_N[mkey] = len(self.fit.params['__mpar__'][mkey][pkey])
        self.update_fixed_parameters()
        self.update_fit_parameters()
        self.update_plot()
        self.xLineEdit.returnPressed.connect(self.xChanged)
        # self.mfitParamTableWidget.cellChanged.connect(self.mfitParamChanged)
        self.fixedParamTableWidget.cellChanged.connect(self.fixedParamChanged)
        self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
        for i in range(self.mfitParamTabWidget.count()):
            mkey = self.mfitParamTabWidget.tabText(i)
            self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)

    def update_fixed_parameters(self):
        try:
            self.fixedParamTableWidget.cellChanged.disconnect()
        except:
            pass
        fpdata=[]        
        for key in self.fit.params.keys():
            if key not in self.fit.fit_params.keys() and key not in self.special_keys and key[:2]!='__':
                fpdata.append((key,str(self.fit.params[key])))
        self.fixedParamData=np.array(fpdata,dtype=[('Params',object),('Value',object)])
        self.fixedParamTableWidget.setData(self.fixedParamData)
        for row in range(self.fixedParamTableWidget.rowCount()):
            self.fixedParamTableWidget.item(row,0).setFlags(Qt.ItemIsEnabled)
            if self.fixedParamTableWidget.item(row, 0).text() in self.fit.params['choices'].keys():
                items=[str(item) for item in self.fit.params['choices'][self.fixedParamTableWidget.item(row,0).text()]]
                combobox=QComboBox()
                combobox.addItems(items)
                self.fixedParamTableWidget.setCellWidget(row,1,combobox)
                index = combobox.findText(str(self.fit.params[self.fixedParamTableWidget.item(row, 0).text()]))
                combobox.setCurrentIndex(index)
                combobox.currentIndexChanged.connect(lambda x: self.fixedParamChanged(row,1))
        self.fixedParamTableWidget.resizeRowsToContents()
        self.fixedParamTableWidget.resizeColumnsToContents()
        self.fixedParamTableWidget.cellChanged.connect(self.fixedParamChanged)


    def update_fit_parameters(self):
        self.update_sfit_parameters()
        # self.update_mfit_parameters()
        self.update_mfit_parameters_new()
        self.sfitParamTableWidget.resizeRowsToContents()
        self.sfitParamTableWidget.resizeColumnsToContents()
        mkeys = list(self.fit.params['__mpar__'].keys())
        if '__mpar__' in self.fit.params.keys() and self.fit.params['__mpar__'] != {}:
            for mkey in mkeys:
                self.mfitParamTableWidget[mkey].resizeRowsToContents()
                self.mfitParamTableWidget[mkey].resizeColumnsToContents()


    def update_sfit_parameters(self):
        try:
            self.sfitParamTableWidget.cellChanged.disconnect()
        except:
            pass
        tpdata=[]
        for key in self.fit.fit_params.keys():
            if key[:2]!='__':
                tpdata.append((key,self.fit.fit_params[key].value,self.fit.fit_params[key].min,
                               self.fit.fit_params[key].max,str(self.fit.fit_params[key].expr),self.fit.fit_params[key].brute_step))
        self.fitParamData=np.array(tpdata,dtype=[('Params',object),('Value',object),('Min',object),('Max',object),
                                                 ('Expr',object),('Brute step',float)])
        self.sfitParamTableWidget.setData(self.fitParamData)
        self.sfitParamTableWidget.setFormat(self.format,column=1)
        for row in range(self.sfitParamTableWidget.rowCount()):
            self.sfitParamTableWidget.item(row,0).setFlags(Qt.ItemIsEnabled)
            par=self.sfitParamTableWidget.item(row,0).text()
            item=self.sfitParamTableWidget.item(row,1)
            item.setFlags(Qt.ItemIsUserCheckable|Qt.ItemIsEnabled|Qt.ItemIsEditable|Qt.ItemIsSelectable)
            if self.fit.fit_params[par].vary==0:
                item.setCheckState(Qt.Unchecked)
            else:
                item.setCheckState(Qt.Checked)
            item.setToolTip((par+' = '+self.format+' \u00B1 '+self.format) % (self.fit.fit_params[par].value, 0.0))
        self.sfitParamTableWidget.resizeRowsToContents()
        self.sfitParamTableWidget.resizeColumnsToContents()
        self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)

    def update_mfit_parameters_new(self):
        self.mfitParamTabWidget.currentChanged.disconnect()
        if '__mpar__' in self.fit.params.keys() and self.fit.params['__mpar__']!={}:
            if len(self.fit.params['__mpar__'])>1:
                self.mfitParamCoupledCheckBox.setEnabled(True)
            # self.mfitParamCoupledCheckBox.setCheckState(Qt.Unchecked)
            self.mfitParamTableWidget = {}
            self.mfitParamData = {}
            mkeys=list(self.fit.params['__mpar__'].keys())
            if self.mfitParamTabWidget.count()>0:
                for i in range(self.mfitParamTabWidget.count()-1,-1,-1):
                    try:
                        self.mfitParamTabWidget.removeTab(i)
                    except:
                        pass
            for mkey in mkeys:
                self.mfitParamTableWidget[mkey] = pg.TableWidget(sortable=False)
                #self.mfitParamTableWidget[mkey].setSelectionBehavior(QAbstractItemView.SelectRows)
                self.mfitParamTableWidget[mkey].cellClicked.connect(self.update_mfitSlider)
                self.mfitParamTableWidget[mkey].cellDoubleClicked.connect(self.mparDoubleClicked)
                self.mfitParamTabWidget.addTab(self.mfitParamTableWidget[mkey],mkey)
                pkeys=list(self.fit.params['__mpar__'][mkey].keys())
                mpar_N=len(self.fit.params['__mpar__'][mkey][pkeys[0]])
                tpdata=[]
                for i in range(mpar_N):
                    temp = []
                    for pkey in pkeys:
                        tkey='__%s_%s_%03d' % (mkey, pkey, i)
                        if tkey in self.fit.fit_params.keys():
                            temp.append(self.fit.fit_params[tkey].value)
                        else:
                            temp.append(self.fit.params['__mpar__'][mkey][pkey][i])
                    tpdata.append(tuple(temp))
                self.mfitParamData[mkey]=np.array(tpdata,dtype=[(pkey,object) for pkey in pkeys])
                self.mfitParamTableWidget[mkey].setData(self.mfitParamData[mkey])
                self.mfitParamTableWidget[mkey].setFormat(self.format)
                for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                    for col in range(self.mfitParamTableWidget[mkey].columnCount()):
                        item = self.mfitParamTableWidget[mkey].item(row, col)
                        if col==0:
                            item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsEditable | Qt.ItemIsSelectable)
                        else:
                            item.setFlags(
                            Qt.ItemIsUserCheckable | Qt.ItemIsEnabled | Qt.ItemIsEditable | Qt.ItemIsSelectable)
                            key = '__%s_%s_%03d' % (mkey, self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text(), row)
                            if self.fit.fit_params[key].vary == 0 or self.fit.fit_params[key].vary==False:
                                item.setCheckState(Qt.Unchecked)
                            else:
                                item.setCheckState(Qt.Checked)
                            item.setToolTip((key + ' = ' + self.format + ' \u00B1 ' + self.format) % (
                            self.fit.fit_params[key].value, 0.0))
                self.mfitParamTableWidget[mkey].resizeRowsToContents()
                self.mfitParamTableWidget[mkey].resizeColumnsToContents()
                self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
            self.add_mpar_button.setEnabled(True)
            self.remove_mpar_button.setEnabled(True)
            self.mfitParamTabChanged(0)
        else:
            self.add_mpar_button.setDisabled(True)
            self.remove_mpar_button.setDisabled(True)
        self.mfitParamTabWidget.currentChanged.connect(self.mfitParamTabChanged)

        
    # def update_mfit_parameters(self):
    #     try:
    #         self.mfitParamTableWidget.cellChanged.disconnect()
    #     except:
    #         pass
    #     if '__mpar__' in self.fit.params.keys() and self.fit.params['__mpar__']!={}:
    #         mpar_keys=list(self.fit.params['__mpar__'].keys())
    #         mpar_N=len(self.fit.params['__mpar__'][mpar_keys[0]])
    #         tpdata=[]
    #         for i in range(mpar_N):
    #             temp=[]
    #             for key in mpar_keys:
    #                 if key in self.fit.fit_params.keys():
    #                     temp.append(self.fit.fit_params['__%s__%03d'%(key,i)].value)
    #                 else:
    #                     temp.append(self.fit.params['__mpar__'][key][i])
    #             tpdata.append(tuple(temp))
    #             #tpdata.append(tuple([self.fit.fit_params['__%s__%03d'%(key,i)].value for key in mpar_keys]))
    #         self.mfitParamData=np.array(tpdata,dtype=[(key,object) for key in mpar_keys])
    #         self.mfitParamTableWidget.setData(self.mfitParamData)
    #         self.mfitParamTableWidget.setFormat(self.format)
    #         self.add_mpar_button.setEnabled(True)
    #         self.remove_mpar_button.setEnabled(True)
    #         for row in range(self.mfitParamTableWidget.rowCount()):
    #             for col in range(1,self.mfitParamTableWidget.columnCount()):
    #                 item=self.mfitParamTableWidget.item(row,col)
    #                 item.setFlags(Qt.ItemIsUserCheckable|Qt.ItemIsEnabled|Qt.ItemIsEditable|Qt.ItemIsSelectable)
    #                 key='__%s__%03d'%(self.mfitParamTableWidget.horizontalHeaderItem(col).text(),row)
    #                 if self.fit.fit_params[key].vary==0:
    #                     item.setCheckState(Qt.Unchecked)
    #                 else:
    #                     item.setCheckState(Qt.Checked)
    #                 item.setToolTip((key + ' = '+self.format+' \u00B1 '+self.format) % (self.fit.fit_params[key].value, 0.0))
    #         self.mfitParamTableWidget.resizeRowsToContents()
    #         self.mfitParamTableWidget.resizeColumnsToContents()
    #     else:
    #         self.add_mpar_button.setDisabled(True)
    #         self.remove_mpar_button.setDisabled(True)
    #         self.mfitParamTableWidget.setData([])
    #     self.mfitParamTableWidget.cellChanged.connect(self.mfitParamChanged)


    def fixedParamChanged(self,row,col):
        try:
            self.fixedParamTableWidget.cellChanged.disconnect()
        except:
            pass
        txt=self.fixedParamTableWidget.item(row,0).text()
        if txt in self.fit.params['choices'].keys():
            self.fixedParamTableWidget.cellWidget(row, 1).currentIndexChanged.disconnect()
            try: # if the parameter is a number
                self.fit.params[txt]=eval(self.fixedParamTableWidget.cellWidget(row,1).currentText())
            except: # if the parameter is a string
                self.fit.params[txt] = str(self.fixedParamTableWidget.cellWidget(row, 1).currentText())
            self.fchanged = False
            self.update_plot()
            self.fixedParamTableWidget.cellWidget(row, 1).currentIndexChanged.connect(lambda x:self.fixedParamChanged(row,1))
        else:
            try: # if the parameter is a number
                val=eval(self.fixedParamTableWidget.item(row,col).text())
            except:  #if the parameter is a string
                val=self.fixedParamTableWidget.item(row,col).text()
            try:
                oldVal=self.fit.params[txt]
                self.fit.params[txt]=val
                self.fchanged = False
                self.update_plot()
            except:
                QMessageBox.warning(self,'Value Error','The value just entered is not seem to be right.\n'+traceback.format_exc(),QMessageBox.Ok)
                self.fixedParamTableWidget.item(row,col).setText(str(oldVal))
        self.fixedParamTableWidget.resizeRowsToContents()
        self.fixedParamTableWidget.resizeColumnsToContents()
        self.update_fit_parameters()
        self.fixedParamTableWidget.cellChanged.connect(self.fixedParamChanged)


        
        
    def sfitParamChanged(self,row,col):
        txt=self.sfitParamTableWidget.item(row,0).text()
        try:
            val=float(self.sfitParamTableWidget.item(row,col).text())
        except:
            val=self.sfitParamTableWidget.item(row,col).text()
        if col==1:
            oldVal=self.fit.fit_params[txt].value
        elif col==2:
            oldVal=self.fit.fit_params[txt].min
        elif col==3:
            oldVal=self.fit.fit_params[txt].vary
        elif col==4:
            oldVal=self.fit.fit_params[txt].expr
        elif col==5:
            oldVal=self.fit.fit_params[txt].brute_step
        if isinstance(val,numbers.Number):
            if col==1:
                if val!=self.fit.fit_params[txt].value:
                    self.fit.params[txt]=val
                    self.fit.fit_params[txt].set(value=val)
                    self.fchanged=False
                    self.sfitParamTableWidget.cellChanged.disconnect()
                    self.sfitParamTableWidget.item(row,col).setText(self.format%val)
                    self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
                    self.update_plot()
            elif col==2:
                self.fit.fit_params[txt].set(min=val)
            elif col==3:
                self.fit.fit_params[txt].set(max=val)
            elif col==5:
                self.fit.fit_params[txt].set(brute_step=val)
        elif isinstance(val,str):
            if col==4:
                pval=self.fit.fit_params[txt].value
                if val == 'None':
                    self.fit.fit_params[txt].set(value=pval,expr = '')
                else:
                    self.fit.fit_params[txt].set(value=pval,expr = val)
                try:
                    self.fit.fit_params[txt].value
                except:
                    self.sfitParamTableWidget.cellChanged.disconnect()
                    QMessageBox.warning(self, 'Expression Error', 'Please enter correct expression using only parameters and constants', QMessageBox.Ok)
                    if oldVal is None:
                        self.fit.fit_params[txt].set(value=pval,expr='')
                    else:
                        self.fit.fit_params[txt].set(value=pval,expr=oldVal)
                    self.sfitParamTableWidget.item(row, col).setText(str(oldVal))
                    self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
            else:
                QMessageBox.warning(self,'Value Error','Please input numbers only',QMessageBox.Ok)
                self.sfitParamTableWidget.cellChanged.disconnect()
                self.sfitParamTableWidget.item(row,col).setText(str(oldVal))
                self.sfitParamTableWidget.cellChanged.connect(self.sfitParamChanged)
        if self.sfitParamTableWidget.item(row,1).checkState()==Qt.Checked:
            self.fit.fit_params[txt].vary=1
        else:
            self.fit.fit_params[txt].vary=0
        if col==1:
            self.sfitParamTableWidget.item(row, 1).setToolTip((txt + ' = '+self.format+'\u00B1 '+self.format) % (self.fit.fit_params[txt].value, 0.0))
        self.update_sfitSlider(row,col)
        self.update_sfit_parameters()
        self.update_mfit_parameters_new()
        self.sfitParamTableWidget.setCurrentCell(row,col)
        self.sfitParamTableWidget.resizeRowsToContents()
        self.sfitParamTableWidget.resizeColumnsToContents()

    def mfitParamChanged_new(self,row,col):
        index=self.mfitParamTabWidget.currentIndex()
        mkey=self.mfitParamTabWidget.tabText(index)
        self.mfitParamTableWidget[mkey].cellChanged.disconnect()
        txt = self.mfitParamTableWidget[mkey].item(row, col).text()
        pkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
        key='__%s_%s_%03d' % (mkey,pkey,row)
        try:
            if col!=0:
                float(txt) # This is for checking the numbers entered to be float or not
                oldval = self.fit.fit_params[key].value
                self.mfitParamTableWidget[mkey].item(row, col).setText(self.format % (float(txt)))
                pchanged=True
                # if float(txt)!=self.fit.fit_params[key].value:
                #     pchanged=True
                #     self.mfitParamTableWidget[mkey].item(row,col).setText(self.format%(float(txt)))
                # else:
                #     self.mfitParamTableWidget[mkey].item(row, col).setText(self.format % (float(txt)))
                #     pchanged=False
                self.fit.fit_params[key].set(value=float(txt))
                if self.mfitParamTableWidget[mkey].item(row,col).checkState()==Qt.Checked:
                    self.fit.fit_params[key].set(vary=1)
                else:
                    self.fit.fit_params[key].set(vary=0)
                self.mfitParamData[mkey][row][col]=float(txt)
                self.fit.fit_params[key].set(value=float(txt))
                self.mfitParamTableWidget[mkey].item(row, col).setToolTip((key + ' = '+self.format+' \u00B1 '+self.format)
                                                                          % (self.fit.fit_params[key].value, 0.0))
                self.fit.params['__mpar__'][mkey][pkey][row]=float(txt)
            else:
                oldval = self.fit.params['__mpar__'][mkey][pkey][row]
                self.fit.params['__mpar__'][mkey][pkey][row] = txt
                self.mfitParamData[mkey][row][col] = txt
                pchanged=True
            self.fchanged=False
            if pchanged:
                try:
                    self.fit.func.output_params={'scaler_parameters': {}}
                    self.update_plot()
                except:
                    QMessageBox.warning(self, 'Value Error', 'The value you entered are not valid!', QMessageBox.Ok)
                    self.mfitParamTableWidget[mkey].item(row, col).setText(oldval)
                    self.fit.params['__mpar__'][mkey][pkey][row] = oldval
                    self.mfitParamData[mkey][row][col]=oldval
            self.update_mfitSlider(row,col)
        except:
            QMessageBox.warning(self,'Value Error', 'Please input numbers only!', QMessageBox.Ok)
            self.mfitParamTableWidget[mkey].item(row,col).setText(str(self.fit.fit_params[key].value))
        self.mfitParamTableWidget[mkey].cellChanged.connect(self.mfitParamChanged_new)
        self.update_fit_parameters()
        self.mfitParamTabWidget.setCurrentIndex(index)
        self.mfitParamTableWidget[mkey].setCurrentCell(row,col)
        item=self.mfitParamTableWidget[mkey].item(row,col)
        item.setSelected(True)
        self.mfitParamTableWidget[mkey].scrollToItem(item)
        self.mfitParamTableWidget[mkey].resizeRowsToContents()
        self.mfitParamTableWidget[mkey].resizeColumnsToContents()


        
    # def mfitParamChanged(self,row,col):
    #     parkey=self.mfitParamTableWidget.horizontalHeaderItem(col).text()
    #     txt=self.mfitParamTableWidget.item(row,col).text()
    #     key = '__%s__%03d' % (parkey, row)
    #     try:
    #         if col!=0:
    #             float(txt) # This is for checking the numbers entered to be float or not
    #             oldval = self.fit.fit_params[key].value
    #             if float(txt)!=self.fit.fit_params[key].value:
    #                 pchanged=True
    #                 self.mfitParamTableWidget.item(row,col).setText(self.format%(float(txt)))
    #             else:
    #                 self.mfitParamTableWidget.item(row, col).setText(self.format % (float(txt)))
    #                 pchanged=False
    #             self.fit.fit_params[key].set(value=float(txt))
    #             if self.mfitParamTableWidget.item(row,col).checkState()==Qt.Checked:
    #                 self.fit.fit_params[key].set(vary=1)
    #             else:
    #                 self.fit.fit_params[key].set(vary=0)
    #             self.mfitParamData[row][col]=float(txt)
    #             self.fit.fit_params[key].set(value=float(txt))
    #             self.mfitParamTableWidget.item(row, col).setToolTip((key + ' = '+self.format+' \u00B1 '+self.format) % (self.fit.fit_params[key].value, 0.0))
    #         else:
    #             oldval = self.fit.params['__mpar__'][parkey][row]
    #             self.fit.params['__mpar__'][parkey][row] = txt
    #             self.mfitParamData[row][col] = txt
    #             pchanged=True
    #         self.fchanged=False
    #         if pchanged:
    #             try:
    #                 self.update_plot()
    #             except:
    #                 QMessageBox.warning(self, 'Value Error', 'The value you entered are not valid!', QMessageBox.Ok)
    #                 self.mfitParamTableWidget.item(row, col).setText(oldval)
    #                 self.fit.params['__mpar__'][parkey][row] = oldval
    #                 self.mfitParamData[row][col]=oldval
    #         self.mfitParamTableWidget.resizeRowsToContents()
    #         self.mfitParamTableWidget.resizeColumnsToContents()
    #         self.update_mfitSlider(row,col)
    #     except:
    #         QMessageBox.warning(self,'Value Error', 'Please input numbers only!', QMessageBox.Ok)
    #         self.mfitParamTableWidget.item(row,col).setText(str(self.fit.fit_params[key].value))
        
            
    def xChanged(self):
        self.xLineEdit.returnPressed.disconnect()
        try:
            x=eval(self.xLineEdit.text())
            #x=np.array(x)
            try:
                self.fit.params['x']=x
                self.fit.set_x(x)
            #self.fit.imin=0
            #self.fit.imax=len(self.fit.x)
            except:
                pass
            self.fchanged=False
            if len(self.funcListWidget.selectedItems())>0:
                try:
                    stime = time.time()
                    self.fit.func.__fit__=False
                    self.fit.evaluate()
                    exectime = time.time() - stime
                except:
                    QMessageBox.warning(self, 'Value error',
                                        'Something wrong with the value of the parameter which you just entered.\n'+traceback.format_exc(),
                                        QMessageBox.Ok)
                    return
                try:
                    self.genParamListWidget.itemSelectionChanged.disconnect()
                except:
                    pass
                self.genParamListWidget.clear()
                self.fit.params['output_params']['scaler_parameters']['Exec-time (sec)'] = exectime
                self.fit.params['output_params']['scaler_parameters']['Chi-Sqr']=self.chisqr
                self.fit.params['output_params']['scaler_parameters']['Red_Chi_Sqr'] = self.red_chisqr
                if len(self.fit.params['output_params']) > 0:
                    for key in self.fit.params['output_params'].keys():
                        if key == 'scaler_parameters':
                            for k in self.fit.params['output_params'][key].keys():
                                self.genParamListWidget.addItem(k + ' : ' + str(self.fit.params['output_params'][key][k]))
                        else:
                            var=[]
                            for k in self.fit.params['output_params'][key].keys():
                                if k!='names' and k!='plotType':
                                    var.append(k)
                            self.genParamListWidget.addItem(str(key) + ' : ' + str(var))
                    if not self.fchanged:
                        for i in range(self.genParamListWidget.count()):
                            item = self.genParamListWidget.item(i)
                            if item.text() in self.gen_param_items:
                                item.setSelected(True)
                    self.plot_extra_param()
                    self.genParamListWidget.itemSelectionChanged.connect(self.plot_extra_param)
                try:
                    pfnames=copy.copy(self.pfnames)
                except:
                    pfnames=[]
                if type(self.fit.x)==dict:
                    for key in self.fit.x.keys():
                        self.plotWidget.add_data(x=self.fit.x[key][self.fit.imin[key]:self.fit.imax[key] + 1], y=self.fit.yfit[key],
                                                 name=self.funcListWidget.currentItem().text()+':'+key, fit=True)
                    pfnames = pfnames + [self.funcListWidget.currentItem().text() + ':' + key for key in
                                             self.fit.x.keys()]
                else:
                    self.plotWidget.add_data(x=self.fit.x[self.fit.imin:self.fit.imax + 1], y=self.fit.yfit,
                                             name=self.funcListWidget.currentItem().text(), fit=True)
                    pfnames = pfnames + [self.funcListWidget.currentItem().text()]

                self.plotWidget.Plot(pfnames)
                # QApplication.processEvents()
                QApplication.processEvents()
        except:
            QMessageBox.warning(self,'Value Error','The value just entered is not seem to be right.\n'+traceback.format_exc(),QMessageBox.Ok)
            self.xLineEdit.setText(self.xline)
        self.xLineEdit.returnPressed.connect(self.xChanged)


    def update_plot(self):
        for row in range(self.fixedParamTableWidget.rowCount()):
            txt=self.fixedParamTableWidget.item(row,0).text()
            if txt in self.fit.params['choices'].keys():
                val = self.fixedParamTableWidget.cellWidget(row, 1).currentText()
            else:
                val=self.fixedParamTableWidget.item(row,1).text()
            try:
                self.fit.params[txt]=eval(val)
            except:
                self.fit.params[txt]=str(val)
        for row in range(self.sfitParamTableWidget.rowCount()):
            txt=self.sfitParamTableWidget.item(row,0).text()
            self.fit.params[txt]=float(self.sfitParamTableWidget.item(row,1).text())
            vary,min,max,expr,bs=self.fit.fit_params[txt].vary,self.fit.fit_params[txt].min,\
                                 self.fit.fit_params[txt].max,self.fit.fit_params[txt].expr,\
                                 self.fit.fit_params[txt].brute_step
            self.fit.fit_params[txt].set(value=float(self.sfitParamTableWidget.item(row,1).text()),vary=vary,min=min,
                                         max=max,expr=expr,brute_step=bs)
        for i in range(self.mfitParamTabWidget.count()):
            mkey=self.mfitParamTabWidget.tabText(i)
            for row in range(self.mfitParamTableWidget[mkey].rowCount()):
                pkey = self.mfitParamTableWidget[mkey].horizontalHeaderItem(0).text()
                txt = self.mfitParamTableWidget[mkey].item(row, 0).text()
                self.fit.params['__mpar__'][mkey][pkey][row] = txt
                for col in range(1,self.mfitParamTableWidget[mkey].columnCount()):
                    pkey=self.mfitParamTableWidget[mkey].horizontalHeaderItem(col).text()
                    txt=self.mfitParamTableWidget[mkey].item(row,col).text()
                    tkey='__%s_%s_%03d'%(mkey,pkey,row)
                    vary,min,max,expr,bs=self.fit.fit_params[tkey].vary,self.fit.fit_params[tkey].min,\
                                         self.fit.fit_params[tkey].max,self.fit.fit_params[tkey].expr,\
                                         self.fit.fit_params[tkey].brute_step
                    self.fit.fit_params['__%s_%s_%03d'%(mkey,pkey,row)].set(value=float(txt),min=min,max=max,vary=vary,expr=expr,brute_step=bs)
        try:
            pfnames=copy.copy(self.pfnames)
        except:
            pfnames=[]
        self.chisqr='None'
        self.red_chisqr='None'
        if len(self.dataListWidget.selectedItems()) > 0:
            if len(self.data[self.sfnames[-1]].keys()) > 1:
                x = {}
                y = {}
                yerr = {}
                for key in self.data[self.sfnames[-1]].keys():
                    x[key] = self.data[self.sfnames[-1]][key]['x']
                    y[key] = self.data[self.sfnames[-1]][key]['y']
                    y[key] = y[key][np.argwhere(x[key] >= self.xmin)[0][0]:np.argwhere(x[key] <= self.xmax)[-1][0]+1]
                    yerr[key] = self.data[self.sfnames[-1]][key]['yerr']
                    yerr[key] = yerr[key][np.argwhere(x[key] >= self.xmin)[0][0]:np.argwhere(x[key] <= self.xmax)[-1][0]+1]
                    x[key] = x[key][np.argwhere(x[key]>=self.xmin)[0][0]:np.argwhere(x[key]<=self.xmax)[-1][0]+1]
            else:
                key = list(self.data[self.sfnames[-1]].keys())[0]
                x = self.data[self.sfnames[-1]][key]['x']
                y = self.data[self.sfnames[-1]][key]['y']
                y = y[np.argwhere(x >= self.xmin)[0][0]:np.argwhere(x <= self.xmax)[-1][0]+1]
                yerr = self.data[self.sfnames[-1]][key]['yerr']
                yerr = yerr[np.argwhere(x >= self.xmin)[0][0]:np.argwhere(x <= self.xmax)[-1][0]+1]
                x = x[np.argwhere(x>=self.xmin)[0][0]:np.argwhere(x<=self.xmax)[-1][0]+1]

        if len(self.funcListWidget.selectedItems())>0:
            try:
                if self.autoCalculate:
                    stime=time.perf_counter()
                    self.fit.func.__fit__=False
                    self.fit.evaluate()
                    ntime=time.perf_counter()
                    exectime=ntime-stime
                else:
                    exectime=0.0
            except:
                QMessageBox.warning(self, 'Evaluation Error', traceback.format_exc(), QMessageBox.Ok)
                self.fit.yfit = self.fit.func.x
                exectime=np.nan
            if len(self.dataListWidget.selectedItems()) > 0:
                self.fit.set_x(x, y=y, yerr=yerr)
                try:
                    residual = self.fit.residual(self.fit.fit_params, self.fitScaleComboBox.currentText())

                    self.chisqr = np.sum(residual ** 2)
                    vary=[self.fit.fit_params[key].vary for key in self.fit.fit_params.keys()]
                    self.red_chisqr=self.chisqr/(len(residual)-np.sum(vary))
                except:
                    QMessageBox.warning(self, 'Evaluation Error', traceback.format_exc(), QMessageBox.Ok)
                    self.chisqr=None
                    self.red_chisqr=None

            try:
                self.genParamListWidget.itemSelectionChanged.disconnect()
            except:
                pass
            self.fitResultTextEdit.clear()
            try:
                self.fitResultTextEdit.append(self.fit_report)
            except:
                self.fitResultTextEdit.clear()
            self.genParamListWidget.clear()
            self.fit.params['output_params']['scaler_parameters']['Exec-time (sec)'] = exectime
            self.fit.params['output_params']['scaler_parameters']['Chi-Sqr'] = self.chisqr
            self.fit.params['output_params']['scaler_parameters']['Red_Chi_Sqr'] = self.red_chisqr
            if len(self.fit.params['output_params'])>0:
                row=0
                for key in self.fit.params['output_params'].keys():
                    if key=='scaler_parameters':
                        for k in self.fit.params['output_params'][key].keys():
                            self.genParamListWidget.addItem(k + ' : ' + str(self.fit.params['output_params'][key][k]))
                            it=self.genParamListWidget.item(row)
                            it.setFlags(it.flags() & ~Qt.ItemIsSelectable)
                            row+=1
                    else:
                        var = []
                        for k in self.fit.params['output_params'][key].keys():
                            if k != 'names' and k != 'plotType':
                                var.append(k)
                        self.genParamListWidget.addItem(
                            str(key) + ' : ' + str(var))
                        row+=1
                if not self.fchanged:
                    for i in range(self.genParamListWidget.count()):
                        item = self.genParamListWidget.item(i)
                        if item.text() in self.gen_param_items:
                            item.setSelected(True)
            self.plot_extra_param()
            self.genParamListWidget.itemSelectionChanged.connect(self.plot_extra_param)
            if type(self.fit.x)==dict:
                for key in self.fit.x.keys():
                    self.plotWidget.add_data(x=self.fit.x[key][self.fit.imin[key]:self.fit.imax[key] + 1], y=self.fit.yfit[key],
                                             name=self.funcListWidget.currentItem().text()+':'+key, fit=True)
                    if len(self.dataListWidget.selectedItems()) > 0:
                        self.fit.params['output_params']['Residuals_%s' % key] = {
                            'x': self.fit.x[key][self.fit.imin[key]:self.fit.imax[key] + 1],
                            'y': (self.fit.y[key][self.fit.imin[key]:self.fit.imax[key] + 1] - self.fit.yfit[key])
                                 / self.fit.yerr[key][self.fit.imin[key]:self.fit.imax[key] + 1]}
                    # else:
                    #     self.fit.params['output_params']['Residuals_%s' % key]={'x':self.fit.x[key][self.fit.imin[key]:self.fit.imax[key] + 1],
                    #                                                             'y':np.zeros_like(self.fit.x[key][self.fit.imin[key]:self.fit.imax[key] + 1])}
                pfnames = pfnames + [self.funcListWidget.currentItem().text() + ':' + key for key in
                                         self.fit.x.keys()]
            else:
                self.plotWidget.add_data(x=self.fit.x[self.fit.imin:self.fit.imax + 1], y=self.fit.yfit,
                                         name=self.funcListWidget.currentItem().text(), fit=True)
                if len(self.dataListWidget.selectedItems()) > 0:
                    self.fit.params['output_params']['Residuals'] = {'x': self.fit.x[self.fit.imin:self.fit.imax + 1],
                                                                     'y': (self.fit.y[
                                                                           self.fit.imin:self.fit.imax + 1] - self.fit.yfit) / self.fit.yerr[
                                                                                                                               self.fit.imin:self.fit.imax + 1]}
                # else:
                #     self.fit.params['output_params']['Residuals'] = {'x': self.fit.x[self.fit.imin:self.fit.imax + 1],
                #                                                      'y':np.zeros_like(self.fit.x[self.fit.imin:self.fit.imax + 1])}
                pfnames=pfnames+[self.funcListWidget.currentItem().text()]
        self.plotWidget.Plot(pfnames)
        # QApplication.processEvents()
        QApplication.processEvents()
        
    def extra_param_doubleClicked(self,item):
        key=item.text().split(':')[0].strip()
        if key in self.fit.params['output_params'].keys():
            if 'x' in self.fit.params['output_params'][key].keys() and 'y' in self.fit.params['output_params'][key].keys():
                x=self.fit.params['output_params'][key]['x']
                y=self.fit.params['output_params'][key]['y']
                if 'meta' in self.fit.params['output_params'][key].keys():
                    meta = self.fit.params['output_params'][key]['meta']
                else:
                    meta = {}
                if 'yerr' in self.fit.params['output_params'][key].keys():
                    yerr=self.fit.params['output_params'][key]['yerr']

                    if 'names' in self.fit.params['output_params'][key].keys():
                        meta['col_names']=self.fit.params['output_params'][key]['names']
                        data = {'data': pd.DataFrame(list(zip(x, y, yerr)), columns=self.fit.params['output_params'][key]['names']),
                            'meta': meta}
                    else:
                        meta['col_names']=['x', 'y', 'yerr']
                        data = {'data': pd.DataFrame(list(zip(x, y, yerr)), columns=['x', 'y', 'yerr']),
                            'meta': meta}

                else:
                    if 'names' in self.fit.params['output_params'][key].keys():
                        meta['col_names'] = self.fit.params['output_params'][key]['names']
                        data = {'data': pd.DataFrame(list(zip(x, y)), columns=self.fit.params['output_params'][key]['names']),
                                'meta': meta}
                    else:
                        meta['col_names'] = ['x', 'y', 'yerr']
                        data = {'data': pd.DataFrame(list(zip(x, y)), columns=['x', 'y']),
                                'meta': meta}
                data_dlg = Data_Dialog(data=data, parent=self, expressions={},
                                       plotIndex=None, colors=None)
                data_dlg.setModal(True)
                data_dlg.closePushButton.setText('Cancel')
                data_dlg.tabWidget.setCurrentIndex(0)
                data_dlg.dataFileLineEdit.setText('None')
                data_dlg.exec_()


    def plot_extra_param(self):
        """
        """
        fdata=[]
        for item in self.genParamListWidget.selectedItems():
            txt,axes=item.text().split(':')
            txt=txt.strip()
            axes=eval(axes)
            if type(axes)==list:
                if len(axes)>=2:
                    x=self.fit.params['output_params'][txt][axes[0]]
                    y=self.fit.params['output_params'][txt][axes[1]]
                    try:
                        yerr=self.fit.params['output_params'][txt][axes[2]]
                    except:
                        yerr=None
                    self.extra_param_1DplotWidget.add_data(x=x,y=y,yerr=yerr,name=txt,fit=True)
                    if 'names' in self.fit.params['output_params'][txt]:
                        self.extra_param_1DplotWidget.setXLabel(self.fit.params['output_params'][txt]['names'][0],fontsize=5)
                        self.extra_param_1DplotWidget.setYLabel(self.fit.params['output_params'][txt]['names'][1],fontsize=5)
                    else:
                        self.extra_param_1DplotWidget.setXLabel('x',fontsize=5)
                        self.extra_param_1DplotWidget.setYLabel('y',fontsize=5)
                    if 'plotType' in self.fit.params['output_params'][txt]:
                        if self.fit.params['output_params'][txt]['plotType']=='step':
                            self.extra_param_1DplotWidget.data[txt].opts['stepMode']='left'
                    fdata.append(txt)
        self.extra_param_1DplotWidget.Plot(fdata)
        self.gen_param_items=[item.text() for item in self.genParamListWidget.selectedItems()]
        # QApplication.processEvents()
        QApplication.processEvents()

if __name__=='__main__':
    # QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    # QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)
    os.environ["QT_AUTO_SCREEN_SCALE_FACTOR"] = "1"
    # app = QApplication(sys.argv)
    # try:
    #     # app.setAttribute(Qt.AA_EnableHighDpiScaling)
    #     app.setHighDpiScaleFactorRoundingPolicy(Qt.HighDpiScaleFactorRoundingPolicy.PassThrough)
    # except:
    #     pass
    QApplication.setAttribute(Qt.AA_EnableHighDpiScaling)
    QApplication.setAttribute(Qt.AA_UseHighDpiPixmaps)
    app = QApplication(sys.argv)
    w=XModFit()
    w.setWindowTitle('XModFit')
    resolution = QDesktopWidget().screenGeometry()
    w.setGeometry(0, 0, resolution.width() - 100, resolution.height() - 100)
    w.move(int(resolution.width() / 2) - int(w.frameSize().width() / 2),
              int(resolution.height() / 2) - int(w.frameSize().height() / 2))
    try:
        fname = sys.argv[1]
        w.addData(fnames=[fname])
    except:
        pass
    try:
        pname=sys.argv[2]
        w.loadParameters(fname=pname)
    except:
        pass
    w.showMaximized()
    # w.show()
    sys.exit(app.exec_())
        
        
        
    
